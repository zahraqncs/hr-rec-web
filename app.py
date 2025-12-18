#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
QNCS HR Management System (Web) — Single File (Revised V5 + Nationality & Dynamic Offers)
Roles: admin, hr, requestor

REVISION V5 NOTES:
1. SCREENING: Removed 'Nationality Type'. 'Nationality' is now a specific dropdown + 'Other'.
2. OFFERS: Dynamic Offer Template selection based on Nationality and Location (Head Office vs Site).
   Supports: offer_saudi, offer_philippine_ho/site, offer_foreign_ho/site.
3. INCLUDES: Fixed CV access, Meeting Links, and Delete functionality from V4.

HOW TO RUN
1) pip install flask pandas openpyxl python-dateutil python-docx
2) python gpt2.py
3) Open http://127.0.0.1:5000

DEFAULT LOGIN
- admin / admin
"""

from __future__ import annotations
import os, io, time, shutil, uuid, csv, mimetypes, json
from datetime import datetime, date, timezone, timedelta
from typing import Dict, Any, List, Tuple, Optional

from flask import (
    Flask, request, redirect, url_for, render_template_string,
    send_file, flash, Response, session
)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash

import pandas as pd
from openpyxl import Workbook, load_workbook
from openpyxl.drawing.image import Image as XLImage

# Try importing docx for Word export
try:
    from docx import Document
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# -----------------------------
# CONFIG
# -----------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
EXCEL_PATH = os.path.join(BASE_DIR, "hr_demo_v2.xlsx")
ATTACH_DIR = os.path.join(BASE_DIR, "Candidates")

# Offer Templates (Filename references)
TPL_SAUDI = "offer_saudi.xlsx"
TPL_PHIL_HO = "offer_philippine_ho.xlsx"
TPL_PHIL_SITE = "offer_philippine_site.xlsx"
TPL_FOREIGN_HO = "offer_foreign_ho.xlsx"
TPL_FOREIGN_SITE = "offer_foreign_site.xlsx"

DEFAULT_OFFER_SHEET = "Sheet1"
LOGO_PATH = os.path.join(BASE_DIR, "qncslogo.png")
OFFER_HEADER_PATH = os.path.join(BASE_DIR, "qncs-header.png")
USERS_PATH = os.path.join(BASE_DIR, "users.json")

OWNER_NAME = "Zahra Aljanabi"

YESNO = ["Yes","No","Other"]
MARITAL = ["Single","Married","Divorced","Widowed","Other"]
IQAMA = ["Valid","Transferable","Expired","N/A (Saudi National)","N/A","Other"]
NOTICE = ["Immediate","1 week","2 weeks","1 month","2 months","3 months","Other"]
EDUCATION = ["High School","Diploma","Bachelor's Degree","Master's Degree","PhD","Other"]
REQ_ACTIONS = ["None", "Schedule Interview", "Proceed to Offer", "Hold", "Reject", "Other"]

# Specific Nationalities List
NATIONALITIES = [
    "Saudi", "Chinese", "Ghanaian", "Lebanese", "Ugandan", "Sudanese", 
    "Yemeni", "Egyptian", "Syrian", "Tunisian", "Palestinian", "Turkish", 
    "Jordanian", "Indian", "Bangladeshi", "Nepalese", "Pakistani", 
    "Greek", "Italian", "Spanish", "Filipino", "Other"
]
INTERVIEW_MODES = ["Online", "Onsite"]

CAND_STATUS = [
    "Screening",
    "Shortlist",
    "Interview",
    "Second Interview",
    "Offer Issued",
    "Offer Accepted",
    "Rejected",
    "On Hold",
    "Back Up",
    "Other",
]

ROLES = ("admin", "hr", "requestor")

THEME = {"main":"#4d85ce","accent":"#d4af37","bg":"#f7faf8","muted":"#a9b4f1"}

STATUS_CLASS_MAP = {
    "screening":"pill screening",
    "shortlist":"pill shortlist",
    "interview":"pill interview",
    "second interview":"pill interview",
    "offer issued":"pill offerissued",
    "offer accepted":"pill offeraccepted",
    "rejected":"pill rejected",
    "on hold":"pill onhold",
    "back up":"pill onhold",
    "other":"pill other",
}

IMPORT_REQUIRED = ["Candidate Name","Role Interviewed For"]
IMPORT_OPTIONAL = [
    "Candidate Email",
    "Phone Number",
    "Total Experience",
    "Relevant Domain Experience",
    "Current Organization",
    "Current Role/Title",
    "Previous Organizations/Roles",
    "Screening Notes",
    "Highest Education",
    "DOB",
    "Marital Status",
    "Family Status (if Married)",
    "Children – Number & Age",
    "Current Location",
    "Desired Location",
    "Nationality",
    "Iqama Status",
    "Profession in Iqama",
    "Current Compensation",
    "Expected Compensation",
    "Notice Period",
    "Ever Interviewed by the client before? (Yes/No)",
    "Recorded By",
    "Gov ID / Iqama / Passport #"
]
IMPORT_COLUMNS = IMPORT_REQUIRED + IMPORT_OPTIONAL

# -----------------------------
# FS & EXCEL HELPERS
# -----------------------------
def ensure_dirs(path: str) -> None:
    os.makedirs(path, exist_ok=True)
ensure_dirs(ATTACH_DIR)

def gen_candidate_id() -> str:
    today = datetime.now().strftime("%Y%m%d")
    unique_suffix = uuid.uuid4().hex[:6].upper()
    return f"CAND-{today}-{unique_suffix}"

def ymd_ok(s: Any) -> bool:
    try:
        datetime.strptime(str(s), "%Y-%m-%d"); return True
    except Exception:
        return False

def normalize_choice(val: Any, allowed: List[str]) -> str:
    v = str(val or "").strip(); low = v.lower()
    for a in allowed:
        if low == a.lower(): return a
    if allowed == YESNO:
        if low in ["y","yes","true","1"]: return "Yes"
        if low in ["n","no","false","0","", "none"]: return "No"
    if allowed == NOTICE:
        mapv = {"immediate":"Immediate","1w":"1 week","2w":"2 weeks","1m":"1 month","2m":"2 months","3m":"3 months"}
        return mapv.get(low, v)
    return v

def candidate_root(name: str, cid: str) -> str:
    safe_name = "".join([c for c in str(name) if c.isalnum() or c in (" ","_","-")]).strip().replace(" ","_")
    return os.path.join(ATTACH_DIR, f"{safe_name}_{cid}")

def candidate_attach_dir(name: str, cid: str) -> str:
    d = os.path.join(candidate_root(name, cid), "Attachments")
    ensure_dirs(d)
    return d

def folder_display_name(name: str, cid: str) -> str:
    return os.path.basename(candidate_root(name, cid))

def _excel_read(sheet: str) -> pd.DataFrame:
    try:
        return pd.read_excel(EXCEL_PATH, sheet_name=sheet)
    except Exception:
        return pd.DataFrame()

def _excel_write(df: pd.DataFrame, sheet: str, retries=3, delay=0.7) -> bool:
    last = None
    for _ in range(retries):
        try:
            with pd.ExcelWriter(EXCEL_PATH, mode="a", if_sheet_exists="replace", engine="openpyxl") as w:
                df.to_excel(w, sheet_name=sheet, index=False)
            return True
        except PermissionError as e:
            last = e; time.sleep(delay)
        except Exception as e:
            last = e; break
    return False

def _generate_offer_doc(cand_id: str, cand_name: str, data: Dict[str, Any], template_filename: str, mapping: Dict[str, str]):
    """
    Generates offer using a specific template and specific mapping.
    """
    safe_name = "".join([c for c in str(cand_name) if c.isalnum() or c in (" ","_","-")]).strip().replace(" ","_")
    cand_folder = os.path.join("Candidates", f"{safe_name}_{cand_id}")

    if not os.path.exists(cand_folder):
        cand_folder = candidate_root(cand_name, cand_id)
        if not os.path.exists(cand_folder):
             raise Exception(f"Candidate folder not found: {cand_folder}")

    tpl_path = os.path.join(BASE_DIR, template_filename)
    if not os.path.exists(tpl_path):
        raise Exception(f"Template not found: {template_filename}. Please upload it to the app folder.")

    wb = load_workbook(tpl_path)
    # Use active sheet or Sheet1
    if DEFAULT_OFFER_SHEET in wb.sheetnames:
        ws = wb[DEFAULT_OFFER_SHEET]
    else:
        ws = wb.active

    # Write data based on mapping
    for key, cell in mapping.items():
        val = data.get(key, "")
        if val is None: val = ""
        ws[cell] = str(val).strip()

    # Try to add logo if it fits standard "A1"
    if os.path.exists(OFFER_HEADER_PATH):
        try:
            img = XLImage(OFFER_HEADER_PATH)
            img.anchor = "A1"
            ws.add_image(img)
        except:
            pass # Ignore if sheet structure doesn't allow

    out_name = f"{cand_id}-offer.xlsx"
    out_path = os.path.join(cand_folder, out_name)
    wb.save(out_path)
    return out_path, None

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set borders on a table cell — works even if template has no styles.

    Example:
    set_cell_border(cell, top=("single", "000000", "1"))
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:tcBorders'
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)

            edge_tag = f"w:{edge}"
            edge_elem = element.find(qn(edge_tag))
            if edge_elem is None:
                edge_elem = OxmlElement(edge_tag)
                element.append(edge_elem)

            edge_elem.set(qn('w:val'), edge_data[0])    # border type
            edge_elem.set(qn('w:sz'),  edge_data[2])    # size
            edge_elem.set(qn('w:color'), edge_data[1])  # color


def _generate_word_spec(cand_id: str, cand_name: str, data: Dict[str,Any]) -> str:
    """
    Generates Screening Word using SF.docx template,
    applies grid borders manually, smaller font,
    saves inside candidate folder.
    """
    if not HAS_DOCX:
        return None

    template_path = os.path.join(BASE_DIR, "SF.docx")

    # Load template (or fallback)
    doc = Document(template_path) if os.path.exists(template_path) else Document()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn


    def force_times_new_roman(doc):
        styles = doc.styles

        for style_name in ["Normal", "Table Normal"]:
            try:
                style = styles[style_name]
            except KeyError:
                continue

            font = style.font
            font.name = "Times New Roman"
            font.size = Pt(10)

            rFonts = style._element.rPr.rFonts
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:eastAsia"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")

    doc = Document(template_path) if os.path.exists(template_path) else Document()

    force_times_new_roman(doc)



    # -------------------------------------

    
    # Candidate folder
    cdir = candidate_root(cand_name, cand_id)
    ensure_dirs(cdir)

    safe = "".join([c for c in cand_name if c.isalnum() or c in ("_","-")]).replace(" ","_")
    out_path = os.path.join(cdir, f"Screening_{safe}_{cand_id}.docx")

    # --------------------------
    # Heading (safe formatting)
    # --------------------------
    head = doc.add_paragraph()
    run = head.add_run(f"Screening Profile: {cand_name}")
    run.bold = True
    run.font.size = Pt(14)
    head.alignment = 1

    # --------------------------
    # Screened fields
    # --------------------------
    fields = [
        ("Candidate ID", cand_id),
        ("Candidate Name", data.get("Candidate Name","")),
        ("Role Interviewed For", data.get("Role Interviewed For","")),
        ("Candidate Email", data.get("Candidate Email","")),
        ("Phone Number", data.get("Phone Number","")),
        ("Highest Education", data.get("Highest Education","")),
        ("DOB", data.get("DOB","")),
        ("Marital Status", data.get("Marital Status","")),
        ("Family Status (if Married)", data.get("Family Status (if Married)","")),
        ("Current Location", data.get("Current Location","")),
        ("Desired Location", data.get("Desired Location","")),
        ("Nationality", data.get("Nationality","")),
        ("Iqama Status", data.get("Iqama Status","")),
        ("Profession in Iqama", data.get("Profession in Iqama","")),
        ("Current Compensation", data.get("Current Compensation","")),
        ("Expected Compensation", data.get("Expected Compensation","")),
        ("Notice Period", data.get("Notice Period","")),
        ("Interviewed Before?", data.get("Ever Interviewed by the client before? (Yes/No)","")),
        ("Recorded By", data.get("Recorded By","")),
        ("Gov ID / Passport", data.get("Gov ID / Iqama / Passport #","")),
        ("Requestor Username", data.get("Requestor Username","")),
        ("Screening Notes", data.get("Screening Notes","")),
    ]

    # --------------------------
    # Create table (no style used)
    # --------------------------
    table = doc.add_table(rows=1, cols=2)

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"

    # Header format
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)

        # Apply borders
        set_cell_border(
            cell,
            top=("single", "000000", "4"),
            left=("single", "000000", "4"),
            bottom=("single", "000000", "4"),
            right=("single", "000000", "4")
        )

    # Data rows
    for label, value in fields:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value) if value else "-"

        # Set small font + borders
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

            set_cell_border(
                cell,
                top=("single", "000000", "4"),
                left=("single", "000000", "4"),
                bottom=("single", "000000", "4"),
                right=("single", "000000", "4")
            )

    # Save final file
    doc.save(out_path)
    return out_path


# -----------------------------
# AUTH (users.json)
# -----------------------------
def _load_users() -> Dict[str, Dict[str,Any]]:
    if not os.path.exists(USERS_PATH):
        data = {"users":{
            "admin":{"username":"admin","name":"Administrator","email":"","role":"admin","password_hash":generate_password_hash("admin")}
        }}
        with open(USERS_PATH, "w", encoding="utf-8") as f: json.dump(data, f, indent=2)
        return data["users"]
    with open(USERS_PATH, "r", encoding="utf-8") as f:
        j = json.load(f) or {}
    return j.get("users", {})

def _save_users(u: Dict[str, Dict[str,Any]]) -> None:
    with open(USERS_PATH, "w", encoding="utf-8") as f:
        json.dump({"users": u}, f, indent=2)

def current_user() -> Optional[Dict[str,Any]]:
    uname = session.get("u")
    if not uname: return None
    users = _load_users()
    return users.get(uname)

def require_login():
    if not current_user():
        flash("Please sign in.", "error"); return False
    return True

def require_role(*roles):
    u = current_user()
    if not u:
        flash("Please sign in.", "error"); return False
    if u.get("role") not in roles:
        flash("Access denied.", "error"); return False
    return True

# -----------------------------
# APP + JINJA
# -----------------------------
app = Flask(__name__)
app.secret_key = os.environ.get("APP_SECRET", "dev-key")

def status_class(s: Any) -> str:
    key = str(s or "").strip().lower()
    return STATUS_CLASS_MAP.get(key, STATUS_CLASS_MAP["other"])
app.jinja_env.globals["status_class"] = status_class

def render_page(tpl: str, **ctx):
    return render_template_string(
        tpl,
        THEME=THEME,
        USER=current_user(),
        LOGO_EXISTS=os.path.exists(LOGO_PATH),
        OWNER_NAME=OWNER_NAME,
        **ctx
    )

# -----------------------------
# HTML BASE
# -----------------------------
BASE_HTML = """
<!doctype html>
<html lang="en" dir="auto">
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width,initial-scale=1"/>
  <title>QNCS HR Management System</title>
  <style>
    *, *::before, *::after { box-sizing: border-box; }
    :root{
      --main:{{ THEME.main }};--accent:{{ THEME.accent }};--muted:{{ THEME.muted }};--bg:{{ THEME.bg }};
      --ok-bg:#e7f8ef;--ok-border:#16a34a;--ok-text:#065f46;
      --err-bg:#fdecec;--err-border:#dc2626;--err-text:#7f1d1d;
    }
    html,body{margin:0;background:var(--bg);color:#162018}
    body{font-family: system-ui,-apple-system,Segoe UI,Roboto,Ubuntu,Cantarell,Noto Sans,sans-serif}
    header{background:var(--main);color:#fff;padding:12px 16px}
    .hdr{display:flex;align-items:center;gap:12px;justify-content:space-between}
    .brand{display:flex;align-items:center;gap:10px;font-weight:900;letter-spacing:.3px}
    .brand img{height:28px;width:auto;border-radius:4px;background:#fff;padding:2px}
    .user{font-weight:600}
    nav{display:flex;gap:10px;flex-wrap:wrap;background:#fff;border-bottom:1px solid #e6e8ea;padding:10px 12px;position:sticky;top:0;z-index:5}
    nav a{color:#0e2d1f;background:transparent;padding:10px 12px;border-radius:10px;text-decoration:none;font-weight:700;border:1px solid transparent}
    nav a:hover{border-color:#dfe5e1;background:#f2f6f4}
    nav a.primary{background:var(--main);color:#fff}
    main{padding:16px;max-width:1200px;margin:auto}
    footer{padding:10px 16px;color:#4d6a59}
    .card{background:#fff;border:1px solid #e5e7eb;border-radius:14px;padding:16px;margin-bottom:16px;box-shadow:0 1px 2px rgba(0,0,0,.04)}
    .row{display:grid;grid-template-columns:repeat(auto-fit,minmax(260px,1fr));gap:12px;align-items:start}
    .col{min-width:0}
    label{font-weight:700;color:var(--main);display:block;margin:8px 0 6px}
    input,select,textarea{width:100%;min-height:40px;padding:10px;border:1px solid #d1d5db;border-radius:10px;background:#fff}
    input[disabled], select[disabled], textarea[disabled] { background: #f3f4f6; color: #6b7280; cursor: not-allowed; }
    input[type=date],input[type=time]{min-height:40px}
    textarea{min-height:90px;resize:vertical}
    table{width:100%;border-collapse:collapse}
    th,td{
  padding:10px;
  border-bottom:1px solid #dbe3f8; /* soft blue border */
}

    th{
  background:var(--main);
  color:#fff;
  position:sticky;
  top:0;
}

    .btn{display:inline-block;padding:8px 10px;border-radius:10px;border:none;cursor:pointer;font-weight:800;line-height:1;text-decoration:none}
    .btn-sm{padding:6px 8px;border-radius:8px;font-weight:700}
    .btn-primary{background:var(--main);color:#fff}
    .btn-accent{background:var(--accent);color:#0e0e0e}
    .btn-ghost{
  background:#e9efff;     /* soft blue */
  color:#2c3e66;
}
.action-buttons {
    display: flex;
    flex-direction: row;
    gap: 8px;
    align-items: center;
    justify-content: flex-start;
}

.action-buttons a {
    white-space: nowrap;
}

    .muted{color:var(--muted);font-style:italic}
    .pill{border-radius:999px;padding:2px 10px;font-size:12px;border:1px solid transparent;display:inline-block}
    .pill.screening{background:#eaf3ff;color:#0b4aa2;border-color:#86a8e7}
    .pill.shortlist{background:#fff7e6;color:#875a07;border-color:#ffd27a}
    .pill.interview{background:#e6fff5;color:#046c4e;border-color:#86e7c4}
    .pill.offerissued{background:#e6fffb;color:#0b766e;border-color:#8fe3dc}
    .pill.offeraccepted{background:#e7f8ef;color:#065f46;border-color:#16a34a}
    .pill.rejected{background:#fdecec;color:#7f1d1d;border-color:#dc2626}
    .pill.onhold{background:#f3f4f6;color:#374151;border-color:#cbd5e1}
    .pill.other{background:#eee;color:#333;border-color:#ddd}
    .alert{border-left:4px solid;padding:10px 12px;border-radius:10px;margin-bottom:12px}
    .alert.ok{background:var(--ok-bg);border-color:var(--ok-border);color:var(--ok-text)}
    .alert.err{background:var(--err-bg);border-color:var(--err-border);color:var(--err-text)}
    .scroll{overflow:auto;max-height:420px}
    .nowrap{white-space:nowrap}
    .actions{display:flex;gap:8px;flex-wrap:wrap;align-items:center}
    .kpi-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(180px,1fr));gap:12px}
    .kpi{background:#fff;border:1px solid #e6e8ea;border-radius:14px;padding:14px}
    .kpi .v{font-size:26px;font-weight:900;color:#0b3f28}
    .kpi .l{color:#4d6a59}
    .no-underline a{text-decoration:none}
    .right{margin-left:auto}
  </style>
  <script>
    // Ctrl+S quick-save
    window.addEventListener('keydown', e => {
      if ((e.ctrlKey || e.metaKey) && e.key === 's') {
        const f = document.querySelector('form');
        if (f) { e.preventDefault(); f.requestSubmit(); }
      }
    });
  </script>
</head>
<body class="no-underline">
<header>
  <div class="hdr">
    <div class="brand">
      {% if LOGO_EXISTS %}
        <img src="{{ url_for('serve_logo') }}" alt="logo"/>
      {% endif %}
      <div>QNCS HR Management System</div>
    </div>
    <div class="user">
      {% if USER %}
        {{ USER.name or USER.username }} ({{ USER.role }}) • <a style="color:#fff;text-decoration:underline" href="{{ url_for('logout') }}">Logout</a>
      {% else %}
        <a style="color:#fff;text-decoration:underline" href="{{ url_for('login') }}">Login</a>
      {% endif %}
    </div>
  </div>
</header>

<nav>
  {% if USER %}
    <a href="{{ url_for('home') }}">Home</a>
    {% if USER.role in ['admin','hr'] %}
      <a href="{{ url_for('screening') }}">Screening</a>
      <a href="{{ url_for('interviews') }}">Interviews</a>
      <a href="{{ url_for('candidates') }}">Candidates Management</a>
      <a href="{{ url_for('offers') }}">Offers</a>
    {% elif USER.role == 'requestor' %}
      <a href="{{ url_for('screening') }}">Screening (View)</a>
      <a href="{{ url_for('interviews') }}">Interviews (List)</a>
      <a href="{{ url_for('candidates') }}">Candidates Management</a>
    {% endif %}
    {% if USER.role == 'admin' %}
      <a class="primary" href="{{ url_for('users_admin') }}">Users</a>
    {% endif %}
  {% endif %}
</nav>

<main>
  {% with msgs = get_flashed_messages(with_categories=True) %}
    {% if msgs %}
      {% for cat,msg in msgs %}
        <div class="alert {{ 'ok' if cat=='success' else 'err' }}">{{ msg }}</div>
      {% endfor %}
    {% endif %}
  {% endwith %}
  {% block content %}{% endblock %}
</main>

<footer>© {{ now().year }} {{ OWNER_NAME }}</footer>
</body>
</html>
"""

# ---------- LOGIN ----------
LOGIN_HTML = """
{% extends "base.html" %}
{% block content %}
<div class="card" style="max-width:520px;margin:auto">
  <h3 style="margin-top:0">Sign in</h3>
  <form method="post">
    <div class="row">
      <div class="col"><label>Username</label><input name="username" autofocus required></div>
      <div class="col"><label>Password</label><input type="password" name="password" required></div>
    </div>
    <button class="btn btn-primary btn-sm" style="margin-top:10px">Login</button>
  </form>
  <div class="muted" style="margin-top:8px">Default: admin / admin (change it in Users).</div>
</div>
{% endblock %}
"""

# ---------- HOME (Dashboard) ----------
HOME_HTML = """
{% extends "base.html" %}
{% block content %}
<div class="card">
  <h2 style="margin:0 0 10px">Dashboard</h2>
  <div class="kpi-grid">
    <div class="kpi"><div class="v">{{ totals.total }}</div><div class="l">Total Candidates</div></div>
    <div class="kpi"><div class="v">{{ totals.shortlisted }}</div><div class="l">Shortlist</div></div>
    <div class="kpi"><div class="v">{{ totals.interview }}</div><div class="l">Interview</div></div>
    <div class="kpi"><div class="v">{{ totals.offer_issued }}</div><div class="l">Offer Issued</div></div>
    <div class="kpi"><div class="v">{{ totals.offer_accepted }}</div><div class="l">Offer Accepted</div></div>
    <div class="kpi"><div class="v">{{ totals.on_hold }}</div><div class="l">On Hold</div></div>
    <div class="kpi"><div class="v">{{ totals.rejected }}</div><div class="l">Rejected</div></div>
  </div>
</div>

<div class="card">
  <h3 style="margin:0 0 8px">By Status</h3>
  <div class="scroll" style="max-height:360px">
    <table>
      <thead><tr><th>Status</th><th>Count</th></tr></thead>
      <tbody>
        {% for s in status_rows %}
        <tr>
          <td><span class="{{ status_class(s.label) }}">{{ s.label }}</span></td>
          <td>{{ s.count }}</td>
        </tr>
        {% endfor %}
      </tbody>
    </table>
  </div>
</div>

<div class="card">
  <h3 style="margin:0 0 8px">Recent Candidates</h3>
  <div class="scroll" style="max-height:360px">
    <table>
      <thead><tr>
        <th class="nowrap">Candidate ID</th><th>Name</th><th>Role</th><th>Status</th><th class="nowrap">Last Updated</th><th>Actions</th>
      </tr></thead>
      <tbody>
        {% for r in recent %}
        <tr>
          <td>{{ r.get('Candidate ID') }}</td>
          <td>{{ r.get('Candidate Name') }}</td>
          <td>{{ r.get('Role') }}</td>
          <td><span class="{{ status_class(r.get('Status')) }}">{{ r.get('Status') }}</span></td>
          <td class="nowrap">{{ r.get('Last Updated') }}</td>
          <td class="actions">

            {% if USER.role in ['admin','hr'] %}

              <a class="btn btn-ghost btn-sm" href="{{ url_for('candidate_detail', cand_id=r.get('Candidate ID')) }}">Open</a>
            {% else %}

              {% if r.get('Requestor Username','') == USER.username %}
                <a class="btn btn-ghost btn-sm" href="{{ url_for('candidate_detail', cand_id=r.get('Candidate ID')) }}">Open</a>
              {% endif %}
            {% endif %}
          </td>
        </tr>
        {% endfor %}
        {% if recent|length == 0 %}
        <tr><td colspan="6" class="muted">No recent records.</td></tr>
        {% endif %}
      </tbody>
    </table>
  </div>
</div>
{% endblock %}
"""

# ---------- USERS (Admin)
USERS_HTML = """
{% extends "base.html" %}
{% block content %}
<div class="card">
  <h3 style="margin-top:0">Users</h3>
  <form method="post" action="{{ url_for('users_create') }}">
    <div class="row">
      <div class="col"><label>Username</label><input name="username" required></div>
      <div class="col"><label>Name</label><input name="name"></div>
      <div class="col"><label>Email</label><input name="email"></div>
      <div class="col"><label>Role</label>
        <select name="role">
          {% for r in ROLES %}<option>{{ r }}</option>{% endfor %}
        </select>
      </div>
      <div class="col"><label>Password</label><input type="password" name="password" required></div>
    </div>
    <button class="btn btn-primary btn-sm" style="margin-top:10px">Add User</button>
  </form>
</div>

<div class="card">
  <h3 style="margin-top:0">All Users</h3>
  <div class="scroll">
    <table>
      <thead><tr><th>Username</th><th>Name</th><th>Email</th><th>Role</th><th>Set Password</th><th>Actions</th></tr></thead>
      <tbody>
        {% for u in users %}
        <tr>
          <form method="post" action="{{ url_for('users_update', username=u.username) }}">
            <td class="nowrap">{{ u.username }}</td>
            <td><input name="name" value="{{ u.name }}"></td>
            <td><input name="email" value="{{ u.email }}"></td>
            <td>
              <select name="role">
                {% for r in ROLES %}<option {% if u.role==r %}selected{% endif %}>{{ r }}</option>{% endfor %}
              </select>
            </td>
            <td><input name="password" placeholder="leave blank to keep"></td>
            <td class="actions">
              <button class="btn btn-primary btn-sm">Save</button>
              {% if USER.username != u.username %}
                <a class="btn btn-ghost btn-sm" href="{{ url_for('users_delete', username=u.username) }}" onclick="return confirm('Delete user?')">Delete</a>
              {% endif %}
            </td>
          </form>
        </tr>
        {% endfor %}
      </tbody>
    </table>
  </div>
</div>
{% endblock %}
"""

# ---------- SCREENING ----------
SCREENING_HTML = """
{% extends "base.html" %}
{% block content %}
{% set is_req = (USER.role == 'requestor') %}


<div class="card">
  <h3 style="margin-top:0">{% if is_req %}View Screening Info (Assigned to You){% else %}Load/Edit Existing Screening{% endif %}</h3>
  <form method="get">
    <div class="row">
      <div class="col">
        <label>Candidate ID & Name</label>
        <select name="pick">
          <option value="">-- select --</option>
          {% for item in picker %}
          <option value="{{ item }}" {% if item == selected %}selected{% endif %}>{{ item }}</option>
          {% endfor %}
        </select>
      </div>
      <div class="col" style="align-self:end">
        <button class="btn btn-primary btn-sm" formaction="{{ url_for('screening_load') }}">Load Info</button>
        {% if not is_req %}
        <a class="btn btn-ghost btn-sm" href="{{ url_for('screening') }}">Start New</a>
        {% endif %}
      </div>
    </div>
  </form>
</div>

<form class="card" method="post" enctype="multipart/form-data" action="{{ url_for('screening_save') }}">
  <div class="row">
    <div class="col"><label>Candidate Name</label><input name="Candidate Name" value="{{ form.get('Candidate Name','') }}" required {% if is_req %}disabled{% endif %}></div>
    <div class="col"><label>Role Interviewed For</label><input name="Role Interviewed For" value="{{ form.get('Role Interviewed For','') }}" required {% if is_req %}disabled{% endif %}></div>
    <div class="col"><label>Candidate Email</label><input name="Candidate Email" value="{{ form.get('Candidate Email','') }}" {% if is_req %}disabled{% endif %}></div>
  </div>

  <div class="row">
    <div class="col"><label>Phone Number</label><input name="Phone Number" value="{{ form.get('Phone Number','') }}" {% if is_req %}disabled{% endif %}></div>
    <div class="col"><label>Total Experience</label><input name="Total Experience" value="{{ form.get('Total Experience','') }}" {% if is_req %}disabled{% endif %}></div>
    <div class="col"><label>Relevant Domain Experience</label><input name="Relevant Domain Experience" value="{{ form.get('Relevant Domain Experience','') }}" {% if is_req %}disabled{% endif %}></div>
  </div>

  <div class="row">
    <div class="col"><label>Current Organization</label><input name="Current Organization" value="{{ form.get('Current Organization','') }}" {% if is_req %}disabled{% endif %}></div>
    <div class="col"><label>Current Role/Title</label><input name="Current Role/Title" value="{{ form.get('Current Role/Title','') }}" {% if is_req %}disabled{% endif %}></div>
    <div class="col"><label>Previous Organizations/Roles</label><input name="Previous Organizations/Roles" value="{{ form.get('Previous Organizations/Roles','') }}" {% if is_req %}disabled{% endif %}></div>
  </div>

  <div class="row">
    <div class="col"><label>Screening Notes</label><textarea name="Screening Notes" {% if is_req %}disabled{% endif %}>{{ form.get('Screening Notes','') }}</textarea></div>
  </div>

  <div class="row">
    <div class="col"><label>Highest Education</label>
      <select name="Highest Education" {% if is_req %}disabled{% endif %}>
        <option value=""></option>
        {% for o in EDUCATION %}<option value="{{o}}" {% if form.get('Highest Education','')==o %}selected{% endif %}>{{o}}</option>{% endfor %}
      </select>
    </div>
    <div class="col">
  <label>DOB</label>
  <input type="date" name="DOB" value="{{ form.get('DOB','') }}" {% if is_req %}disabled{% endif %}>
  <label>Age</label>
<input type="number" id="ageField" name="Age" 
       value="{{ form.get('Age','') }}" 
       placeholder="Auto">

<button type="button" id="calcAgeBtn" style="margin-top:5px;">Calculate Age</button>


</div>

    <div class="col"><label>Marital Status</label>
      <select name="Marital Status" {% if is_req %}disabled{% endif %}>
        <option value=""></option>
        {% for o in MARITAL %}<option value="{{o}}" {% if form.get('Marital Status','')==o %}selected{% endif %}>{{o}}</option>{% endfor %}
      </select>
    </div>
   <div class="col" id="family-status-wrapper" style="display:none;">
  <label>Family Status (If Married)</label>
  <select name="Family Status (if Married)" {% if is_req %}disabled{% endif %}>
    <option value=""></option>
    <option value="Residential (Iqama)"
      {% if form.get('Family Status (if Married)','') == 'Residential (Iqama)' %}selected{% endif %}>
      Residential (Iqama)
    </option>
    <option value="Visit Visa"
      {% if form.get('Family Status (if Married)','') == 'Visit Visa' %}selected{% endif %}>
      Visit Visa
    </option>
    <option value="N/A"
      {% if form.get('Family Status (if Married)','') == 'N/A' %}selected{% endif %}>
      N/A
    </option>
  </select>
</div>


  </div>

  <div class="row">
    <div class="col"><label>Children – Number & Age</label><input name="Children – Number & Age" value="{{ form.get('Children – Number & Age','') }}" {% if is_req %}disabled{% endif %}></div>
    <div class="col"><label>Current Location</label><input name="Current Location" value="{{ form.get('Current Location','') }}" {% if is_req %}disabled{% endif %}></div>
    <div class="col"><label>Desired Location</label><input name="Desired Location" value="{{ form.get('Desired Location','') }}" {% if is_req %}disabled{% endif %}></div>
  </div>

  <div class="row">
    <div class="col">
      <label>Nationality</label>
      <select name="nationality_select" id="natSelect" onchange="toggleNatOther()" {% if is_req %}disabled{% endif %}>
          <option value="">-- select --</option>
          {% for n in NATIONALITIES %}
            <option value="{{ n }}" 
              {% if form.get('Nationality','') == n or (n=='Other' and form.get('Nationality','') not in NATIONALITIES and form.get('Nationality','')) %}selected{% endif %}
            >{{ n }}</option>
          {% endfor %}
      </select>
      <input type="text" name="nationality_other" id="natOther" placeholder="Type nationality..." 
             style="display:none; margin-top:5px;" 
             value="{% if form.get('Nationality','') not in NATIONALITIES %}{{ form.get('Nationality','') }}{% endif %}" 
             {% if is_req %}disabled{% endif %}>
    </div>
    <div class="col"><label>Iqama Status</label>
      <select name="Iqama Status" {% if is_req %}disabled{% endif %}>
        <option value=""></option>
        {% for o in IQAMA %}<option value="{{o}}" {% if form.get('Iqama Status','')==o %}selected{% endif %}>{{o}}</option>{% endfor %}
      </select>
    </div>
  </div>

  <div class="row">
    <div class="col"><label>Profession in Iqama</label><input name="Profession in Iqama" value="{{ form.get('Profession in Iqama','') }}" {% if is_req %}disabled{% endif %}></div>
    <div class="col"><label>Current Compensation (Package)</label><input name="Current Compensation" value="{{ form.get('Current Compensation','') }}" {% if is_req %}disabled{% endif %}></div>
    <div class="col"><label>Expected Compensation (Package)</label><input name="Expected Compensation" value="{{ form.get('Expected Compensation','') }}" {% if is_req %}disabled{% endif %}></div>
  </div>

  <div class="row">
    <div class="col"><label>Notice Period</label>
      <select name="Notice Period" {% if is_req %}disabled{% endif %}>
        <option value=""></option>
        {% for o in NOTICE %}<option value="{{o}}" {% if form.get('Notice Period','')==o %}selected{% endif %}>{{o}}</option>{% endfor %}
      </select>
    </div>
    <div class="col"><label>Ever Interviewed by the client before? (Yes/No)</label>
      <select name="Ever Interviewed by the client before? (Yes/No)" {% if is_req %}disabled{% endif %}>
        <option value=""></option>
        {% for o in YESNO %}<option value="{{o}}" {% if form.get('Ever Interviewed by the client before? (Yes/No)','')==o %}selected{% endif %}>{{o}}</option>{% endfor %}
      </select>
    </div>
    <div class="col"><label>Recorded By</label><input name="Recorded By" value="{{ form.get('Recorded By','') }}" {% if is_req %}disabled{% endif %}></div>
  </div>

  <div class="row">
    <div class="col"><label>Gov ID / Iqama / Passport #</label><input name="Gov ID / Iqama / Passport #" value="{{ form.get('Gov ID / Iqama / Passport #','') }}" {% if is_req %}disabled{% endif %}></div>
    <div class="col">
      <label>CV/Resume Attachment (optional)</label>
      <div class="row">
        <div class="col"><input name="cv_file" type="file" {% if is_req %}disabled{% endif %}></div>
        <div class="col"><input name="cv_existing" placeholder="Existing CV path" value="{{ form.get('CV File Path','') }}" readonly></div>
      </div>
      {% if not is_req %}<div class="muted">You can upload now or later; both are fine.</div>{% endif %}
    </div>
    <div class="col">
      <label>Requestor (Account)</label>
      <select name="Requestor Username" {% if is_req %}disabled{% endif %}>
        <option value=""></option>
        {% for r in REQUESTORS %}<option value="{{r.username}}" {% if form.get('Requestor Username','')==r.username %}selected{% endif %}>{{ r.name or r.username }}</option>{% endfor %}
      </select>
      {% if not is_req %}<div class="muted">HR picks who requested this candidate.</div>{% endif %}
    </div>
  </div>

  <input type="hidden" name="Candidate ID" value="{{ form.get('Candidate ID','') }}"/>
  <div style="margin-top:12px" class="actions">
    {% if not is_req %}
      <button class="btn btn-primary btn-sm">Save Screening</button>
    {% endif %}
    {% if form.get('Candidate ID') %}
      <a class="btn btn-accent btn-sm" href="{{ url_for('screening_download', cand_id=form.get('Candidate ID')) }}">Download Word Spec</a>
      {% if not is_req %}
        <a class="btn btn-ghost btn-sm" href="{{ url_for('screening_delete', cand_id=form.get('Candidate ID')) }}" onclick="return confirm('Delete screening + candidate?')">Delete</a>
      {% endif %}
    {% endif %}
  </div>
</form>

<script>
  function toggleFamilyStatus() {
    var ms = document.querySelector('select[name="Marital Status"]');
    var wrap = document.getElementById('family-status-wrapper');
    if (!ms || !wrap) return;
    if (ms.value === 'Married') {
      wrap.style.display = '';
    } else {
      wrap.style.display = 'none';
      var sel = wrap.querySelector('select[name="Family Status (if Married)"]');
      if (sel) sel.value = '';
    }
  }

  function toggleNatOther() {
    var sel = document.getElementById('natSelect');
    var oth = document.getElementById('natOther');
    if (!sel || !oth) return;
    if (sel.value === 'Other') {
        oth.style.display = 'block';
    } else {
        oth.style.display = 'none';
    }
  }

// =============================
// AGE AUTO CALC + BUTTON CALC
// =============================
document.addEventListener('DOMContentLoaded', function () {

    const dobInput = document.querySelector('input[name="DOB"]');
    const ageField = document.getElementById('ageField');
    const calcBtn = document.getElementById('calcAgeBtn');

    function calcAge() {
        if (!dobInput || !dobInput.value) return;
        const dob = new Date(dobInput.value);
        const today = new Date();

        let age = today.getFullYear() - dob.getFullYear();
        const month = today.getMonth() - dob.getMonth();

        if (month < 0 || (month === 0 && today.getDate() < dob.getDate())) {
            age--;
        }

        if (age >= 0) {
            ageField.value = age;
        }
    }

    // Auto calc when DOB changes
    if (dobInput) {
        dobInput.addEventListener('change', calcAge);
    }

    // Manual button calc
    if (calcBtn) {
        calcBtn.addEventListener('click', calcAge);
    }
});
document.addEventListener('DOMContentLoaded', function () {
  const ms = document.querySelector('select[name="Marital Status"]');
  if (ms) {
    ms.addEventListener('change', toggleFamilyStatus);
    toggleFamilyStatus(); // 🔥 auto-apply on load
  }
});


</script>
{% endblock %}
"""

INTERVIEWS_HTML = """
{% extends "base.html" %}
{% block content %}

{% if USER.role in ['admin','hr'] %}
<div class="card">
  <h3 style="margin-top:0">Interview Form</h3>

  <form method="post" action="{{ url_for('interviews') }}">
    <div class="row">
      <div class="col">
        <label>Candidate</label>
        <select name="pick" id="candPick" required>
          <option value="">-- select --</option>
          {% for combo in combos %}
            <option value="{{combo}}" {% if edit_combo == combo %}selected{% endif %}>{{combo}}</option>
          {% endfor %}
        </select>
      </div>
      <div class="col">
        <label>Email</label>
        <input id="candEmail" readonly placeholder="auto" 
               value="{% if edit_iv %}{{edit_iv.get('Email','')}}{% endif %}">
      </div>
      <div class="col">
        <label>Position</label>
        <input id="candPos" name="Position" placeholder="auto"
               value="{% if edit_iv %}{{edit_iv.get('Position','')}}{% endif %}">
      </div>
    </div>

    <div class="row">
      <div class="col">
        <label>Date</label>
        <input type="date" name="Interview Date" 
               value="{% if edit_iv %}{{ edit_iv.get('Interview Date','') }}{% endif %}" required>
      </div>
      <div class="col">
        <label>Time</label>
        <input type="time" name="Interview Time" 
               value="{% if edit_iv %}{{ edit_iv.get('Interview Time','') }}{% endif %}" required>
      </div>
      <div class="col">
        <label>Mode</label>
        <select name="Mode" id="interviewMode" onchange="updateInterviewFields()" required>
          <option value="">-- select --</option>
          {% for mode in INTERVIEW_MODES %}
            <option value="{{mode}}" {% if edit_iv and edit_iv.get('Mode')==mode %}selected{% endif %}>{{mode}}</option>
          {% endfor %}
        </select>
      </div>
    </div>

    <div class="row">
      <div class="col">
        <label>Location</label>
        <input name="Location" 
value="{% if edit_iv %}{{edit_iv.get('Location/Link','')}}{% endif %}">
      </div>
      <div class="col" id="meetingLinkCol">
        <label>Meeting Link</label>
        <input name="Meeting Link" value="{% if edit_iv %}{{edit_iv.get('Meeting Link','')}}{% endif %}">
      </div>
      <div class="col">
        <label>Interviewer</label>
        <input name="Interviewer Name" value="{% if edit_iv %}{{edit_iv.get('Interviewer','')}}{% endif %}">
      </div>
    </div>

    <input type="hidden" name="is_second" value="{% if second_mode %}1{% else %}0{% endif %}">

    <button class="btn btn-primary btn-sm" style="margin-top:12px;">Save Interview</button>
  </form>
</div>
{% endif %}

<div class="card">
  <h3 style="margin-top:0">Scheduled Interviews</h3>
  <div class="scroll">
    <table>
      <thead>
        <tr>
          <th>Candidate</th>
          <th>Position</th>
          <th>Date & Time</th>
          <th>Mode</th>
          <th>Location / Link</th>
          <th>Interviewer</th>
          <th>Status</th>
          {% if USER.role in ['admin','hr'] %}
            <th>Actions</th>
          {% endif %}
        </tr>
      </thead>

      <tbody>
      {% for iv in existing_interviews %}
        {% set st = iv.get('Status','') %}
        {% set date = iv.get('Interview Date') %}
        {% set time = iv.get('Interview Time') %}
        {% set is_past = (date < now().strftime("%Y-%m-%d")) %}

        <tr>

          <td>{{ iv.get('Candidate Name') }}</td>
          <td>{{ iv.get('Position') }}</td>

          <td {% if is_past %}style="color:red; font-weight:800"{% endif %}>
            {{ date }} {{ time }}
            {% if is_past %}
              <span style="color:white; background:red; padding:2px 6px; border-radius:6px; margin-left:6px; font-size:11px;">
                PASSED
              </span>
            {% endif %}
          </td>

          <td>{{ iv.get('Mode') }}</td>

          <td>
            {% if iv.get('Mode') == 'Onsite' %}
                <a href="{{ iv.get('Location/Link') }}" target="_blank">Location</a>
            {% else %}
                {% if iv.get('Meeting Link') %}
                    <a href="{{ iv.get('Meeting Link') }}" target="_blank">Join</a>
                {% endif %}
            {% endif %}
          </td>

          <td>{{ iv.get('Interviewer') }}</td>

          <td>
            {% if st == 'First Interview' %}
              <span class="pill interview pill-sm">First Interview</span>
            {% elif st == 'First Interview Completed' %}
              <span class="pill offeraccepted pill-sm">First Interview Completed</span>
            {% elif st == 'Second Interview' %}
              <span class="pill interview pill-sm">Second Interview</span>
            {% elif st == 'Second Interview Completed' %}
              <span class="pill offeraccepted pill-sm">Second Interview Completed</span>
            {% else %}
              <span class="pill other pill-sm">{{ st or 'N/A' }}</span>
            {% endif %}
          </td>

          {% if USER.role in ['admin','hr'] %}
          <td class="actions">

            {% if st == 'First Interview' %}
              <a class="btn btn-primary btn-sm"
                 href="{{ url_for('interview_done_first', idx=loop.index0) }}">Done</a>

            {% elif st == 'First Interview Completed' %}
              <a class="btn btn-accent btn-sm"
                 href="{{ url_for('interview_request_second', idx=loop.index0) }}">Second Interview</a>
              <a class="btn btn-ghost btn-sm"
                 href="{{ url_for('interview_undo_first', idx=loop.index0) }}">Undo</a>

            {% elif st == 'Second Interview' %}
              <a class="btn btn-primary btn-sm"
                 href="{{ url_for('interview_done_second', idx=loop.index0) }}">Done</a>

            {% elif st == 'Second Interview Completed' %}
              <a class="btn btn-ghost btn-sm"
                 href="{{ url_for('interview_undo_second', idx=loop.index0) }}">Undo</a>
            {% endif %}

            <a class="btn btn-ghost btn-sm"
               href="{{ url_for('interviews', edit_cid=iv.get('Candidate ID')) }}">Edit</a>

            <a class="btn btn-ghost btn-sm"
               href="{{ url_for('interview_delete', idx=loop.index0) }}"
               onclick="return confirm('Delete?')">Delete</a>

            {# -------- COMPOSE EMAIL -------- #}

            {% set name = iv.get('Candidate Name') %}
            {% set email = iv.get('Email') %}
            {% set mode = iv.get('Mode') %}
            {% set loc = iv.get('Location/Link') %}
            {% set link = iv.get('Meeting Link') %}

            {% if mode == 'Online' %}
              {% set body =
                "Dear " ~ name ~ ",%0D%0A%0D%0A"
                ~ "We are pleased to invite you to an online interview at Qudrat National Commercial Services Company.%0D%0A%0D%0A"
                ~ "Interview Details:%0D%0A"
                ~ "• Date: " ~ date ~ "%0D%0A"
                ~ "• Time (KSA Time): " ~ time ~ "%0D%0A"
                ~ "• Interview Format: Microsoft Teams%0D%0A"
                ~ "• Teams Link: " ~ link ~ "%0D%0A%0D%0A"
                ~ "Please ensure:%0D%0A• Stable internet%0D%0A• Camera on%0D%0A• Quiet place%0D%0A%0D%0A"
                ~ "We look forward to speaking with you."
              %}
            {% else %}
              {% set body =
                "Dear " ~ name ~ ",%0D%0A%0D%0A"
                ~ "We are pleased to invite you to a second interview at Qudrat National Commercial Commercial Company.%0D%0A%0D%0A"
                ~ "Interview Details:%0D%0A"
                ~ "• Date: " ~ date ~ "%0D%0A"
                ~ "• Time: " ~ time ~ "%0D%0A"
                ~ "• Location: " ~ loc ~ "%0D%0A"
                ~ "2nd Floor, Building 862, Alqashlah, Al Khobar 34232.%0D%0A%0D%0A"
                ~ "Please bring your updated CV.%0D%0A%0D%0A"
                ~ "We look forward to meeting you."
              %}
            {% endif %}

      {% if iv.get('ICS Path') %}
<a class="btn btn-sm"
   style="background:#A9D3FF; color:black;"
   href="{{ url_for('open_appointment', path=iv.get('ICS Path')) }}">
   Create Outlook Appointment
</a>
{% endif %}



          </td>
          {% endif %}
        </tr>
      {% endfor %}
      </tbody>
    </table>
  </div>
</div>

<script>
const CAND_MAP = {{cand_map_json|safe}};

function syncCandInfo(){
  const pick = document.getElementById('candPick');
  const email = document.getElementById('candEmail');
  const pos = document.getElementById('candPos');
  const info = CAND_MAP[pick.value] || {};
  if(email && !email.value) email.value = info.email || "";
  if(pos && !pos.value) pos.value = info.role || "";
}

function updateInterviewFields(){
  const mode = document.getElementById('interviewMode').value;
  const loc = document.querySelector('input[name="Location"]');
  const link = document.querySelector('input[name="Meeting Link"]');
  const col = document.getElementById('meetingLinkCol');

  if(mode === 'Onsite'){
    loc.value = "https://maps.app.goo.gl/WLVM69QkbntDG2Tq7";
    loc.readOnly = true;
    link.value = "";
    link.disabled = true;
    col.style.display = "none";
  } else {
    loc.readOnly = false;
    link.disabled = false;
    col.style.display = "";
  }
}

document.addEventListener("DOMContentLoaded",()=>{
  const pick = document.getElementById('candPick');
  if(pick) pick.addEventListener("change",syncCandInfo);
  syncCandInfo();
  updateInterviewFields();
});
</script>

{% endblock %}
"""

# ---------- OFFERS ----------
OFFERS_HTML = """
{% extends "base.html" %}
{% block content %}
<div class="card">
  <h3 style="margin-top:0">Offer Generation</h3>
  <form method="post" action="{{ url_for('offer_load') }}">
    <div class="row">
      <div class="col">
        <label>Select Candidate (Name [ID])</label>
        <select name="pick" required>
          <option value="">-- select --</option>
          {% for combo in combos %}<option value="{{combo}}" {% if combo==selected %}selected{% endif %}>{{combo}}</option>{% endfor %}
        </select>
      </div>
      <div class="col" style="align-self:end">
        <button class="btn btn-primary btn-sm">Load Candidate</button>
      </div>
    </div>
  </form>
</div>

<form class="card" method="post" action="{{ url_for('offer_generate') }}">

  <div class="row">
    <div class="col"><label>Position</label><input name="Role Interviewed For" value="{{ data.get('Role Interviewed For','') }}"></div>
    <div class="col"><label>Candidate Name</label><input name="Candidate Name" value="{{ data.get('Candidate Name','') }}"></div>
    <div class="col"><label>Gov ID / Iqama / Passport #</label><input name="Gov ID / Iqama / Passport #" value="{{ data.get('Gov ID / Iqama / Passport #','') }}"></div>
  </div>
  <div class="row">
    <div class="col"><label>Nationality</label><input name="Nationality" value="{{ data.get('Nationality','') }}"></div>
    <div class="col"><label>Candidate Email</label><input name="Candidate Email" value="{{ data.get('Candidate Email','') }}"></div>
    <div class="col"><label>Offer Issue Date</label><input name="Offer Issue Date" value="{{ data.get('Offer Issue Date', today) }}"></div>
  </div>
  
  <div class="row">
     <div class="col">
        <label>Work Location Type</label>
        <select name="Location Type">
           <option value="Head Office">Head Office</option>
           <option value="Site">Site</option>
        </select>
     </div>
     <div class="col"><label>Marital Status</label><input name="Marital Status" value="{{ data.get('Marital Status','') }}"></div>
  </div>

  <h4 style="margin:16px 0 8px; border-bottom:1px solid #eee;">Financials</h4>
  
  <div class="row">
    <div class="col"><label>Basic Salary</label><input name="Basic Salary" value="{{ data.get('Basic Salary','') }}" placeholder="e.g. 5000"></div>
    <div class="col"><label>Accommodation Allowance</label><input name="Accommodation Allowance" value="{{ data.get('Accommodation Allowance','') }}" placeholder="e.g. 1250"></div>
    <div class="col"><label>Transportation Allowance</label><input name="Transportation Allowance" value="{{ data.get('Transportation Allowance','') }}" placeholder="e.g. 500"></div>
  </div>
  
  <div class="row">
    <div class="col"><label>Monthly Fixed Allowance</label><input name="Monthly Fixed Allowance" value="{{ data.get('Monthly Fixed Allowance','') }}"></div>
    <div class="col"><label>Other Monthly Allowance</label><input name="Other Monthly Allowance" value="{{ data.get('Other Monthly Allowance','') }}"></div>
    <div class="col"><label>Air Ticket Amount </label><input name="Air Ticket" value="{{ data.get('Air Ticket','') }}"></div>
  </div>

  <input type="hidden" name="selected" value="{{ selected or '' }}">
  
  <div style="margin-top:12px" class="actions">
    <button class="btn btn-primary btn-sm">Generate & Save Offer (.xlsx)</button>
    <div class="muted">
      System auto-selects template (Saudi, Foreign HO/Site, Philippine HO/Site) based on Nationality & Location.
      <br>Total Package is calculated automatically from filled fields.
    </div>
  </div>
</form>
{% endblock %}
"""
CANDIDATES_HTML = """
{% extends "base.html" %}
{% block content %}
<div class="card">
  <h3 style="margin-top:0">Candidates Management</h3>
<form method="get">
  <div class="row">
    <div class="col">
      <label>Filter by Name</label>
      <input name="name" value="{{ q.get('name','') }}">
    </div>

    <div class="col">
      <label>Filter by Role</label>
      <input name="role" value="{{ q.get('role','') }}">
    </div>

    <div class="col">
      <label>Filter by ID / Iqama</label>
      <input name="search_id" value="{{ q.get('search_id','') }}">
    </div>

    <div class="col">
      <label>Status</label>
      <select name="status">
        <option value="All" {% if q.get('status','All')=='All' %}selected{% endif %}>All</option>
        {% for s in statuses %}
        <option {% if q.get('status','All')==s %}selected{% endif %}>{{s}}</option>
        {% endfor %}
      </select>
    </div>

    <div class="col" style="align-self:end">
      <button class="btn btn-primary btn-sm">Apply</button>
    </div>
  </div>
</form>


  <div class="scroll" style="margin-top:10px">
    <table>
      <thead><tr>
        <th class="nowrap">Candidate ID</th>
        <th>Gov ID / Iqama</th>
        <th>Candidate Name</th>
        <th>Role</th>
        <th>Nationality</th>

        <th>Status</th><th>HR Owner</th><th class="nowrap">Last Updated</th><th>Actions</th>
      </tr></thead>
      <tbody>
        {% for r in rows %}
        <tr>
          <td>{{ r.get('Candidate ID') }}</td>
          <td>{{ r.get('Gov ID / Iqama / Passport #') }}</td>
          <td>{{ r.get('Candidate Name') }}</td>
          <td>{{ r.get('Role') }}</td>

          <td>{{r.get('Nationality')}}</td>
          <td><span class="{{ status_class(r.get('Status')) }}">{{r.get('Status')}}</span></td>
          <td>{{r.get('HR Owner')}}</td>
          <td class="nowrap">{{r.get('Last Updated')}}</td>

{% if USER.role in ['admin','hr'] %}
<td>
    <div class="action-buttons">
        <a class="btn btn-ghost btn-sm"
           href="{{ url_for('candidate_detail', cand_id=r.get('Candidate ID')) }}">Open</a>

        <a class="btn btn-ghost btn-sm"
           href="{{ url_for('screening_delete', cand_id=r.get('Candidate ID')) }}"
           onclick="return confirm('Delete candidate + screening?')">Delete</a>
    </div>
</td>


{% elif USER.role == 'requestor' %}
<td class="actions" style="display:flex; gap:6px; align-items:center;">
    {% if r.get('Requestor Username','') == USER.username %}
        <a class="btn btn-ghost btn-sm"
           href="{{ url_for('candidate_detail', cand_id=r.get('Candidate ID')) }}">Open</a>
    {% endif %}
</td>
{% endif %}

        </tr>
        {% endfor %}
      </tbody>
    </table>
  </div>
</div>

{% if cand %}
  {% if USER.role in ['admin','hr'] %}
    <div class="card">
      <h3 style="margin:0 0 10px">Candidate — {{ cand_name }} [{{ cand_id }}]</h3>
      <form method="post" action="{{ url_for('candidate_meta_save', cand_id=cand_id) }}">
        <div class="row">
          <div class="col">
            <label>Status</label>
            <select name="Status">
              {% for s in statuses %}<option value="{{s}}" {% if meta.get('Status')==s %}selected{% endif %}>{{s}}</option>{% endfor %}
            </select>
          </div>
          <div class="col"><label>Next Action</label><input name="Next Action" value="{{ meta.get('Next Action','') }}"></div>
          <div class="col"><label>Notes</label><input name="Notes" value="{{ meta.get('Notes','') }}"></div>
        </div>
        <div class="row" style="margin-top:10px; border-top:1px solid #eee; padding-top:10px;">
            <div class="col">
                <label>Requestor Action</label>
                <input value="{{ meta.get('Requestor Action','') }}" readonly class="muted" style="background:#f8f9fa">
            </div>
            <div class="col">
                <label>Suggested Date</label>
                <input value="{{ meta.get('Suggested Interview Date','') }}" readonly class="muted" style="background:#f8f9fa">
            </div>
             <div class="col">
                <label>Suggested Time</label>
                <input value="{{ meta.get('Suggested Interview Time','') }}" readonly class="muted" style="background:#f8f9fa">
            </div>
             <div class="col">
                <label>Requestor Comments</label>
                <div style="background:#f8f9fa; padding:8px; border-radius:8px; border:1px solid #ddd; min-height:40px;">
                    {{ meta.get('Requestor Comments','') }}
                </div>
            </div>
        </div>
        <div class="actions" style="margin-top:10px; display:flex; gap:6px; align-items:center;">
          <button class="btn btn-primary btn-sm">Save Status & Notes</button>
          <span class="{{ status_class(meta.get('Status')) }}" style="margin-left:8px">{{ meta.get('Status') }}</span>
        </div>
      </form>
    </div>

    <div class="card">
      <h3 style="margin-top:0">Shortlist — {{ cand_name }} [{{ cand_id }}]</h3>
      <form method="post" action="{{ url_for('shortlist_save', cand_id=cand_id) }}" enctype="multipart/form-data">
        <div class="row">
          <div class="col"><label>Add Item</label><input name="new_item" placeholder="e.g., Passport/Iqama Copy"></div>
          <div class="col" style="align-self:end"><button class="btn btn-primary btn-sm">Add</button></div>
        </div>
        <div class="scroll" style="margin-top:10px;max-height:360px">
          <table>
            <thead><tr><th>Item</th><th>Received (Yes/No)</th><th>Notes</th><th>Mapped File</th><th class="center">Actions</th></tr></thead>
            <tbody>
              {% for it in checklist %}
              {% set p = (it.get('Mapped File Path') or '')|string|lower %}
              <tr>
                <td><input name="item_{{loop.index}}" value="{{it['Item']}}"></td>
                <td>
                  <select name="recv_{{loop.index}}">
                    {% for o in YESNO %}<option {% if it['Received (Yes/No)']==o %}selected{% endif %}>{{o}}</option>{% endfor %}
                  </select>
                </td>
                <td><input name="note_{{loop.index}}" value="{{it.get('Notes','')}}"></td>
              <td class="center">
    <div class="action-buttons">

        {% if p not in ['', 'none', 'nan'] %}
            <a class="btn btn-ghost btn-sm"
               href="{{ url_for('open_inline', path=it['Mapped File Path']) }}"
               target="_blank">Open</a>
        {% endif %}

       {% if p in ['', 'none', 'nan'] %}
    <input type="file" name="map_{{loop.index}}">
{% endif %}


        <a class="btn btn-ghost btn-sm"
           href="{{ url_for('shortlist_delete', cand_id=cand_id, idx=loop.index0) }}">
           Delete Item
        </a>

        {% if p not in ['', 'none', 'nan'] %}
           <a class="btn btn-danger btn-sm"
   href="{{ url_for('shortlist_remove_file', cand_id=cand_id, idx=loop.index0) }}"
   onclick="return confirm('Remove attachment only?')">
   Remove File
</a>

        {% endif %}

    </div>
</td>

              </tr>
              {% endfor %}
            </tbody>
          </table>
        </div>
        <input type="hidden" name="rows" value="{{ checklist|length }}">
        <div class="actions" style="margin-top:10px; display:flex; gap:6px;">
            <button class="btn btn-primary btn-sm">Save Checklist</button>
        </div>
      </form>
    </div>

  {% else %}
    <div class="card">
      <h3 style="margin-top:0">Attachments & Documents</h3>
      <div class="scroll">
        <table>
          <thead><tr><th>Document Name</th><th>Action</th></tr></thead>
          <tbody>

            {# CV BLOCK FIXED #}
            {% set cvp = meta.get('CV File Path')|string|lower %}
            {% if cvp not in ['', 'none', 'nan'] %}
            <tr>
                <td>CV / Resume</td>
                <td>
                   <a class="btn btn-ghost btn-sm"
                      href="{{ url_for('open_inline', path=meta['CV File Path']) }}"
                      target="_blank">View</a>
                </td>
            </tr>
            {% endif %}

            {# SHORTLIST DOCUMENTS FIXED #}
            {% for it in checklist %}
              {% set p = it.get('Mapped File Path')|string|lower %}
              {% if p not in ['', 'none', 'nan'] %}
              <tr>
                <td>{{ it.get('Item') }}</td>
                <td>
                    <a class="btn btn-ghost btn-sm"
                       href="{{ url_for('open_inline', path=it['Mapped File Path']) }}"
                       target="_blank">View</a>
                </td>
              </tr>
              {% endif %}
            {% endfor %}

            {% if cvp in ['', 'none', 'nan'] and checklist|selectattr('Mapped File Path')|list|length == 0 %}
               <tr><td colspan="2" class="muted">No documents attached yet.</td></tr>
            {% endif %}
          </tbody>
        </table>
      </div>
    </div>

    <div class="card">
      <h3 style="margin:0 0 10px">Candidate — {{ cand_name }} [{{ cand_id }}]</h3>
      <form method="post" action="{{ url_for('candidate_comment', cand_id=cand_id) }}">
        
        <div class="alert ok">You can suggest interview dates or request actions here. You cannot schedule directly.</div>
        
        <div class="row">
           <div class="col">
              <label>Action Requested</label>
              <select name="Requestor Action">
                  <option value="">-- select --</option>
                  {% for act in req_actions %}
                     <option value="{{ act }}" {% if meta.get('Requestor Action')==act %}selected{% endif %}>{{ act }}</option>
                  {% endfor %}
              </select>
           </div>
           <div class="col">
              <label>Suggested Interview Date</label>
              <input type="date" name="Suggested Interview Date" value="{{ meta.get('Suggested Interview Date','') }}">
           </div>
           <div class="col">
              <label>Suggested Interview Time</label>
              <input type="time" name="Suggested Interview Time" value="{{ meta.get('Suggested Interview Time','') }}">
           </div>
        </div>

        <div class="row" style="margin-top:12px">
          <div class="col"><label>Comment to HR</label><textarea name="Requestor Comments" placeholder="Type your note to HR...">{{ meta.get('Requestor Comments','') }}</textarea></div>
        </div>

        <div class="actions" style="margin-top:10px; display:flex; gap:6px;">
            <button class="btn btn-primary btn-sm">Send Request to HR</button>
        </div>
      </form>
      
      <div style="margin-top:20px; border-top:1px solid #eee; padding-top:10px;">
         <h4>Current Status: <span class="{{ status_class(meta.get('Status')) }}">{{ meta.get('Status') }}</span></h4>
         <p><strong>HR Notes:</strong> {{ meta.get('Notes') or 'No notes visible.' }}</p>
      </div>
    </div>
  {% endif %}
{% endif %}
{% endblock %}
"""


# -----------------------------
# ROUTES
# -----------------------------
@app.context_processor
def inject_now():
    from datetime import datetime as _dt
    return {"now": _dt.now, "NATIONALITIES": NATIONALITIES}

@app.get("/assets/logo")
def serve_logo():
    if not os.path.exists(LOGO_PATH): return "No logo", 404
    return send_file(LOGO_PATH, mimetype=mimetypes.guess_type(LOGO_PATH)[0] or "image/png")

# -------- Auth
@app.route("/login", methods=["GET","POST"])
def login():
    if request.method == "GET":
        return render_page(LOGIN_HTML)
    uname = (request.form.get("username") or "").strip()
    pw    = request.form.get("password") or ""
    users = _load_users()
    u = users.get(uname)
    if not u or not check_password_hash(u["password_hash"], pw):
        flash("Invalid credentials.", "error"); return redirect(url_for("login"))
    session["u"] = u["username"]
    flash("Welcome.", "success")
    return redirect(url_for("home"))

@app.get("/logout")
def logout():
    session.pop("u", None)
    flash("Signed out.", "success")
    return redirect(url_for("login"))

# -------- Users (Admin)
@app.get("/admin/users")
def users_admin():
    if not require_role("admin"): return redirect(url_for("login"))
    users = list(_load_users().values())
    users = sorted(users, key=lambda x: x["username"].lower())
    for u in users:  # ensure keys
        for k in ("username","name","email","role"):
            u.setdefault(k,"")
    return render_page(USERS_HTML, users=users, ROLES=ROLES)

@app.post("/admin/users/create")
def users_create():
    if not require_role("admin"): return redirect(url_for("login"))
    users = _load_users()
    uname = (request.form.get("username") or "").strip()
    if not uname or uname in users:
        flash("Username exists / invalid.", "error"); return redirect(url_for("users_admin"))
    role = request.form.get("role","requestor").strip()
    if role not in ROLES: role = "requestor"
    users[uname] = {
        "username": uname,
        "name": request.form.get("name",""),
        "email": request.form.get("email",""),
        "role": role,
        "password_hash": generate_password_hash(request.form.get("password") or "changeme"),
    }
    _save_users(users)
    flash("User added.", "success")
    return redirect(url_for("users_admin"))

@app.post("/admin/users/update/<username>")
def users_update(username: str):
    if not require_role("admin"): return redirect(url_for("login"))
    users = _load_users()
    u = users.get(username)
    if not u:
        flash("Not found.", "error"); return redirect(url_for("users_admin"))
    u["name"]  = request.form.get("name","")
    u["email"] = request.form.get("email","")
    r = request.form.get("role","requestor")
    u["role"] = r if r in ROLES else u["role"]
    pw = request.form.get("password","").strip()
    if pw: u["password_hash"] = generate_password_hash(pw)
    _save_users(users)
    flash("User updated.", "success")
    return redirect(url_for("users_admin"))

@app.get("/admin/users/delete/<username>")
def users_delete(username: str):
    if not require_role("admin"): return redirect(url_for("login"))
    me = current_user()
    if me and me["username"] == username:
        flash("Cannot delete yourself.", "error"); return redirect(url_for("users_admin"))
    users = _load_users()
    if username in users:
        users.pop(username)
        _save_users(users)
        flash("User deleted.", "success")
    else:
        flash("Not found.", "error")
    return redirect(url_for("users_admin"))

# -------- Home
@app.route("/")
def home():
    if not require_login(): return redirect(url_for("login"))
    totals, status_rows, recent = _dashboard_data()
    # If requestor, filter recent to own candidates only
    u = current_user()
    if u and u.get("role")=="requestor":
        recent = [r for r in recent if (r.get("Requestor Username","")==u["username"])]
    return render_page(HOME_HTML, totals=totals, status_rows=status_rows, recent=recent)

# -------- Screening (HR/Admin/Requestor)
@app.get("/screening")
def screening():
    if not require_role("admin","hr","requestor"): return redirect(url_for("login"))
    
    u = current_user()
    # Pass username to picker to filter list if requestor
    picker_items = _screening_picker(filter_user=(u["username"] if u["role"]=="requestor" else None))

    return render_page(SCREENING_HTML,
        picker=picker_items, selected="", form={},
        YESNO=YESNO, EDUCATION=EDUCATION, MARITAL=MARITAL, IQAMA=IQAMA, NOTICE=NOTICE,
        REQUESTORS=_list_requestors(), NATIONALITIES=NATIONALITIES
    )

@app.get("/screening/load")
def screening_load():
    if not require_role("admin","hr","requestor"): return redirect(url_for("login"))
    pick = request.args.get("pick","").strip()
    if not pick or "[" not in pick:
        flash("Select a screening entry.", "error"); return redirect(url_for("screening"))
    cid = pick.split("[")[-1].strip("]")
    row = _load_screening_row(cid)
    
    # Access check for requestor
    u = current_user()
    if u.get("role") == "requestor":
        if str(row.get("Requestor Username","")) != u["username"]:
             flash("Access denied to this candidate's screening.", "error")
             return redirect(url_for("screening"))

    if not row:
        flash("Screening row not found.", "error"); return redirect(url_for("screening"))
    
    # Re-filter picker
    picker_items = _screening_picker(filter_user=(u["username"] if u["role"]=="requestor" else None))

    return render_page(SCREENING_HTML,
        picker=picker_items, selected=pick, form=row,
        YESNO=YESNO, EDUCATION=EDUCATION, MARITAL=MARITAL, IQAMA=IQAMA, NOTICE=NOTICE,
        REQUESTORS=_list_requestors(), NATIONALITIES=NATIONALITIES
    )

@app.get("/screening/template")
def screening_template():
    if not require_role("admin","hr"): return redirect(url_for("login"))
    output = io.StringIO(); writer = csv.writer(output)
    header = IMPORT_COLUMNS + ["Requestor Username"]
    writer.writerow(header)
    writer.writerow([
        "Jane Doe",                     # Candidate Name
        "Facilities Supervisor",        # Role Interviewed For
        "jane@example.com",             # Candidate Email
        "0500000000",                   # Phone Number
        "7",                            # Total Experience
        "5",                            # Relevant Domain Experience
        "ACME Co.",                     # Current Organization
        "Supervisor",                   # Current Role/Title
        "Company X; Company Y",         # Previous Organizations/Roles
        "Good profile",                 # Screening Notes
        "Bachelor's Degree",            # Highest Education
        "1995-06-20",                   # DOB
        "Single",                       # Marital Status
        "",                             # Family Status (if Married)
        "–",                            # Children – Number & Age
        "Dammam",                       # Current Location
        "Riyadh",                       # Desired Location
        "Saudi",                        # Nationality
        "Valid",                        # Iqama Status
        "Engineer",                     # Profession in Iqama
        "8000",                         # Current Compensation
        "9000",                         # Expected Compensation
        "Immediate",                    # Notice Period
        "No",                           # Ever Interviewed...
        "Zahra",                        # Recorded By
        "1234567890",                   # Gov ID / Iqama / Passport #
        ""                              # Requestor Username
    ])
    csv_bytes = io.BytesIO(output.getvalue().encode("utf-8"))
    return send_file(csv_bytes, mimetype="text/csv", as_attachment=True, download_name="screening_single_template.csv")

@app.post("/screening/import")
def screening_import():
    if not require_role("admin","hr"): return redirect(url_for("login"))
    f = request.files.get("import_file")
    if not f or not f.filename:
        flash("Select an import file.", "error"); return redirect(url_for("screening"))
    ext = os.path.splitext(f.filename)[1].lower()
    try:
        if ext in [".xlsx",".xls"]: df = pd.read_excel(f)
        elif ext == ".csv": df = pd.read_csv(f)
        else:
            flash("Use .xlsx or .csv.", "error"); return redirect(url_for("screening"))
    except Exception:
        flash("File read error.", "error"); return redirect(url_for("screening"))
    df = df.replace({pd.NA:"", None:""}).dropna(how="all")
    if df.empty: flash("No data found.", "error"); return redirect(url_for("screening"))
    row = df.iloc[0].to_dict()
    for col in IMPORT_REQUIRED:
        if not str(row.get(col,"")).strip():
            flash(f'Missing "{col}" in first row.', "error"); return redirect(url_for("screening"))
    cand_name = str(row.get("Candidate Name")).strip()
    role = str(row.get("Role Interviewed For")).strip()
    cand_id = gen_candidate_id()

    norm = {}
    for k in IMPORT_COLUMNS + ["Requestor Username"]:
        v = row.get(k, "")
        if k == "Gov ID / Iqama / Passport #":
            v = str(v).strip().replace(".0","")
        else:
            v = str(v).strip()
        norm[k] = v

    if norm.get("DOB") and not ymd_ok(norm["DOB"]): norm["DOB"] = ""
    norm["Highest Education"] = normalize_choice(norm.get("Highest Education",""), EDUCATION)
    norm["Marital Status"] = normalize_choice(norm.get("Marital Status",""), MARITAL)
    norm["Iqama Status"] = normalize_choice(norm.get("Iqama Status",""), IQAMA)
    norm["Notice Period"] = normalize_choice(norm.get("Notice Period",""), NOTICE)
    norm["Ever Interviewed by the client before? (Yes/No)"] = normalize_choice(norm.get("Ever Interviewed by the client before? (Yes/No)",""), YESNO)

    # Screening_Form write
    sf = _excel_read("Screening_Form")
    required_cols = {
        "Timestamp","Candidate ID","Candidate Name","Role Interviewed For","Candidate Email","Phone Number",
        "Total Experience","Relevant Domain Experience","Current Organization","Current Role/Title",
        "Previous Organizations/Roles","Screening Notes","Highest Education","DOB","Marital Status",
        "Family Status (if Married)","Children – Number & Age","Current Location","Desired Location",
        "Nationality","Iqama Status","Profession in Iqama",
        "Current Compensation","Expected Compensation","Notice Period",
        "Ever Interviewed by the client before? (Yes/No)","Recorded By",
        "Gov ID / Iqama / Passport #","CV File Path","Requestor Username","Age"

    }
    if sf.empty: sf = pd.DataFrame(columns=sorted(list(required_cols)))
    srow = {"Timestamp": pd.Timestamp.now(), "Candidate ID": cand_id, "CV File Path": ""}
    for k in required_cols:
        if k in norm: srow[k] = norm[k]
    for col in sf.columns:
        if col not in srow: srow[col] = ""
    sf = pd.concat([sf, pd.DataFrame([srow])], ignore_index=True)
    if not _excel_write(sf, "Screening_Form"):
        flash("Write error (Screening_Form).", "error"); return redirect(url_for("screening"))

    # Candidates sheet sync
    cd = _excel_read("Candidates")
    base_cols = [
        "Candidate ID","Candidate Name","Role","Nationality","Status","Requestor Assessment",
        "HR Owner","Next Action","CV File Path","Last Updated","Notes",
        "Requestor Username","Requestor Comments", "Requestor Action",
        "Suggested Interview Date", "Suggested Interview Time"
    ]
    if cd.empty: cd = pd.DataFrame(columns=base_cols)
    new_cand = {
        "Candidate ID": cand_id, "Candidate Name": cand_name, "Role": role,
        "Nationality": norm.get("Nationality",""),
        "Status": "Screening", "Requestor Assessment": "Pending",
        "HR Owner": norm.get("Recorded By",""), "Next Action": "Review screening details",
        "CV File Path": "", "Last Updated": datetime.now(), "Notes": "",
        "Requestor Username": norm.get("Requestor Username",""), "Requestor Comments":"",
        "Requestor Action": "", "Suggested Interview Date": "", "Suggested Interview Time": ""
    }
    for col in base_cols:
        if col not in new_cand: new_cand[col] = ""
    cd = pd.concat([cd, pd.DataFrame([new_cand])], ignore_index=True)
    if not _excel_write(cd, "Candidates"):
        flash("Write error (Candidates).", "error"); return redirect(url_for("screening"))

    short_name = folder_display_name(cand_name, cand_id)
    flash(f"Screening saved Candidate ID: {cand_id} (Folder: {short_name})", "success")
    return redirect(url_for("screening_load", pick=f"{cand_name} [{cand_id}]"))

@app.post("/screening/save")
def screening_save():
    if not require_role("admin","hr"): 
        # Requestors cannot save screening edits
        flash("Access denied. Read-only.", "error")
        return redirect(url_for("screening"))
        
    form = request.form.to_dict()

    # Basic fields
    cand_name = (form.get("Candidate Name") or form.get("Name") or "").strip()
    role = (form.get("Role Interviewed For") or form.get("Role") or form.get("Position") or "").strip()
    cv_existing = (form.get("cv_existing") or "").strip()
    cv_upload = request.files.get("cv_file")
    req_username = (form.get("Requestor Username") or "").strip()
    
    # Nationality Logic
    nat_select = form.get("nationality_select","")
    nat_other = form.get("nationality_other","").strip()
    nationality = nat_other if nat_select == "Other" else nat_select

    # Resolve candidate_id
    candidate_id = (form.get("Candidate ID") or form.get("CandidateID") or form.get("ID") or form.get("cid") or "").strip()
    if not candidate_id and cand_name:
        try:
            df_c = _excel_read("Candidates")
            row = df_c[df_c["Candidate Name"].astype(str).str.strip() == cand_name].head(1)
            if not row.empty:
                candidate_id = str(row.iloc[0].get("Candidate ID", "")).strip()
        except Exception:
            pass
    if not candidate_id:
        candidate_id = gen_candidate_id()

    # Validate
    if not cand_name or not role:
        flash("Missing required fields.", "error")
        return redirect(url_for("screening"))

    # Prepare folder / files
    out_dir = candidate_root(cand_name, candidate_id)
    ensure_dirs(out_dir)
    cv_path = ""
    try:
        if cv_upload and getattr(cv_upload, "filename", ""):
            from werkzeug.utils import secure_filename
            dst_dir = candidate_attach_dir(cand_name, candidate_id)
            fn = secure_filename(cv_upload.filename)
            target = os.path.join(dst_dir, fn)
            cv_upload.save(target)
            cv_path = target
        elif cv_existing:
            cv_path = cv_existing.strip()
    except Exception:
        pass

    def _norm(v): 
        return ("" if v is None else str(v)).strip()

    norm = {k: _norm(form.get(k)) for k in [
        "Candidate Name","Role Interviewed For","Candidate Email","Phone Number","Total Experience",
        "Relevant Domain Experience","Current Organization","Current Role/Title",
        "Previous Organizations/Roles","Screening Notes","Highest Education","DOB",
        "Marital Status","Family Status (if Married)","Children – Number & Age",
        "Current Location","Desired Location",
        "Iqama Status","Profession in Iqama",
        "Current Compensation","Expected Compensation","Notice Period",
        "Ever Interviewed by the client before? (Yes/No)","Recorded By","Gov ID / Iqama / Passport #",
        "Requestor Username"
    ]}
    # AGE FIELD (Manual or Auto From UI)
    norm["Age"] = _norm(form.get("Age"))

    
    # Force-clean Gov ID (remove Excel float .0)
    gid = norm.get("Gov ID / Iqama / Passport #", "")
    gid = str(gid).strip().replace(".0", "")
    norm["Gov ID / Iqama / Passport #"] = gid

    norm["Candidate Name"] = cand_name
    norm["Role Interviewed For"] = role
    norm["Nationality"] = nationality
    norm["Requestor Username"] = req_username

    if norm.get("DOB") and not ymd_ok(norm["DOB"]): norm["DOB"] = ""
    norm["Highest Education"] = normalize_choice(norm.get("Highest Education",""), EDUCATION)
    norm["Marital Status"] = normalize_choice(norm.get("Marital Status",""), MARITAL)
    norm["Iqama Status"] = normalize_choice(norm.get("Iqama Status",""), IQAMA)
    norm["Notice Period"] = normalize_choice(norm.get("Notice Period",""), NOTICE)
    norm["Ever Interviewed by the client before? (Yes/No)"] = normalize_choice(norm.get("Ever Interviewed by the client before? (Yes/No)",""), YESNO)

    # Upsert Screening
    sf = _excel_read("Screening_Form")
    required_cols = {
        "Timestamp","Candidate ID","Candidate Name","Role Interviewed For","Candidate Email","Phone Number",
        "Total Experience","Relevant Domain Experience","Current Organization","Current Role/Title",
        "Previous Organizations/Roles","Screening Notes","Highest Education","DOB","Marital Status",
        "Family Status (if Married)","Children – Number & Age","Current Location","Desired Location",
        "Nationality","Iqama Status","Profession in Iqama",
        "Current Compensation","Expected Compensation","Notice Period",
        "Ever Interviewed by the client before? (Yes/No)","Recorded By",
        "Gov ID / Iqama / Passport #","CV File Path","Requestor Username","Age"

    }

    if sf.empty:
        sf = pd.DataFrame(columns=sorted(list(required_cols)))
    for c in required_cols:
        if c not in sf.columns: sf[c] = ""
    mask = (sf["Candidate ID"].astype(str) == str(candidate_id))
    if mask.any():
        sf.loc[mask, "Timestamp"] = pd.Timestamp.now()
        for k in required_cols:
            if k in norm: sf.loc[mask, k] = norm[k]
        sf.loc[mask, "CV File Path"] = cv_path
    else:
        srow = {"Timestamp": pd.Timestamp.now(), "Candidate ID": candidate_id, "CV File Path": cv_path}
        for k in required_cols:
            if k in norm: srow[k] = norm[k]
            elif k not in srow: srow[k] = ""
        sf = pd.concat([sf, pd.DataFrame([srow])], ignore_index=True)
    _excel_write(sf, "Screening_Form")

    # Upsert Candidates
    cd = _excel_read("Candidates")
    base_cols = [
        "Candidate ID","Candidate Name","Role","Nationality","Status","Requestor Assessment",
        "HR Owner","Next Action","CV File Path","Last Updated","Notes",
        "Requestor Username","Requestor Comments", "Requestor Action",
        "Suggested Interview Date", "Suggested Interview Time"
    ]
    if cd.empty: cd = pd.DataFrame(columns=base_cols)
    for c in base_cols:
        if c not in cd.columns: cd[c] = ""

    new_vals = {
        "Candidate Name": cand_name,
        "Role": role,
        "Nationality": nationality,
        "Status": "Screening",
        "Requestor Assessment": cd["Requestor Assessment"].iloc[0] if not cd.empty and "Requestor Assessment" in cd.columns else "Pending",
        "HR Owner": norm.get("Recorded By",""),
        "Next Action": "Review screening details",
        "CV File Path": cv_path,
        "Last Updated": datetime.now(),
        "Notes": "",
        "Requestor Username": req_username,
    }
    gid = norm.get("Gov ID / Iqama / Passport #","").strip().replace(".0","")
    new_vals["Gov ID / Iqama / Passport #"] = gid

    m = (cd["Candidate ID"].astype(str) == str(candidate_id))
    if m.any():
        for k, v in new_vals.items():
            cd.loc[m, k] = v
    else:
        row = {"Candidate ID": candidate_id}
        row.update(new_vals)
        cd = pd.concat([cd, pd.DataFrame([row])], ignore_index=True)
    _excel_write(cd, "Candidates")

    flash("Screening saved.", "success")
    return redirect(url_for("screening_load", pick=f"{cand_name} [{candidate_id}]"))


@app.get("/screening/download/<cand_id>")
def screening_download(cand_id: str):
    # Requestors can download
    if not require_role("admin","hr","requestor"): return redirect(url_for("login"))
    
    row = _load_screening_row(cand_id)
    u = current_user()
    # Requestor guard
    if u.get("role") == "requestor":
         if str(row.get("Requestor Username","")) != u["username"]:
             flash("Access denied.", "error"); return redirect(url_for("screening"))

    if not row:
        flash("Not found.", "error"); return redirect(url_for("screening"))
    
    cname = row.get("Candidate Name","")
    try:
        path = _generate_word_spec(cand_id, cname, row)
        if not path:
            flash("Word export requires python-docx installed.", "error")
            return redirect(url_for("screening_load", pick=f"{cname} [{cand_id}]"))
        return send_file(path, as_attachment=True, download_name=os.path.basename(path))
    except Exception as e:
        flash(f"Export failed: {e}", "error")
        return redirect(url_for("screening"))

@app.get("/screening/delete/<cand_id>")
def screening_delete(cand_id: str):
    if not require_role("admin","hr"): return redirect(url_for("login"))
    sf = _excel_read("Screening_Form")
    if not sf.empty:
        sf = sf[sf["Candidate ID"].astype(str)!=str(cand_id)]
        _excel_write(sf, "Screening_Form")
    cd = _excel_read("Candidates")
    if not cd.empty:
        cd = cd[cd["Candidate ID"].astype(str)!=str(cand_id)]
        _excel_write(cd, "Candidates")
    for entry in os.listdir(ATTACH_DIR):
        if entry.endswith("_"+cand_id):
            try: shutil.rmtree(os.path.join(ATTACH_DIR, entry))
            except Exception: pass
    flash(f"Deleted Candidate ID: {cand_id}", "success")
    return redirect(url_for("screening"))

@app.route("/interviews", methods=["GET","POST"])
def interviews():
    # Allowed for Admin/HR (create/view) and Requestor (view only)
    if not require_role("admin","hr","requestor"):
        return redirect(url_for("login"))

    u = current_user()
    is_req = (u.get("role") == "requestor")

    # Load existing interviews
    iv_df = _excel_read("Interviews")
    existing_interviews = []
    if not iv_df.empty:
        for col in ["Meeting Link", "Location/Link", "Status"]:
            if col not in iv_df.columns:
                iv_df[col] = ""
        if is_req:
            cand_df = _excel_read("Candidates")
            my_cands = set()
            if not cand_df.empty:
                for _, r in cand_df.iterrows():
                    if str(r.get("Requestor Username","")) == u.get("username",""):
                        my_cands.add(str(r.get("Candidate ID","")))
            for _, row in iv_df.iterrows():
                if str(row.get("Candidate ID","")) in my_cands:
                    existing_interviews.append(row.to_dict())
        else:
            existing_interviews = iv_df.to_dict("records")

    # candidate combos
    cand_map = _interview_cand_map()
    combos = _candidate_combo_list()

    edit_cid = request.args.get("edit_cid","").strip()
    second_cid = request.args.get("second_cid","").strip()
    chosen_cid = edit_cid or second_cid

    edit_iv = None
    edit_combo = ""
    second_mode = bool(second_cid)

    if chosen_cid and not iv_df.empty and "Candidate ID" in iv_df.columns:
        rows = iv_df[iv_df["Candidate ID"].astype(str) == chosen_cid]
        if not rows.empty:
            # use the last row for that candidate
            edit_iv = rows.iloc[-1].to_dict()
            cname_edit = str(edit_iv.get("Candidate Name","")).strip()
            edit_combo = f"{cname_edit} [{chosen_cid}]"

    if request.method == "GET":
        last_ics = request.args.get("last_ics","").strip()
        mailto = request.args.get("mailto","").strip()

        return render_page(
            INTERVIEWS_HTML,
            combos=combos,
            cand_map_json=json.dumps(cand_map),
            last_ics=(last_ics if last_ics and os.path.exists(last_ics) else None),
            last_ics_name=os.path.basename(last_ics) if last_ics else None,
            auto_open_url="",
            mailto_url=mailto or None,
            existing_interviews=existing_interviews,
            INTERVIEW_MODES=INTERVIEW_MODES,
            edit_iv=edit_iv,
            edit_combo=edit_combo,
            second_mode=second_mode,
        )

    # POST: create/update interview (Admin/HR only)
    if is_req:
        flash("Requestors cannot create interviews.", "error")
        return redirect(url_for("interviews"))

    form = request.form.to_dict()
    pick = form.get("pick","").strip()
    if not pick or "[" not in pick:
        flash("Select a candidate.", "error")
        return redirect(url_for("interviews"))

    cid = pick.split("[")[-1].strip("]")
    cname = pick.split("[")[0].strip()

    cmap_entry = cand_map.get(pick, {}) or {}
    email = cmap_entry.get("email","")
    position = form.get("Position","").strip() or cmap_entry.get("role","")

    date_str = form.get("Interview Date","").strip()
    time_str = form.get("Interview Time","").strip()
    if not date_str or not time_str:
        flash("Pick date and time.", "error")
        return redirect(url_for("interviews"))

    try:
        start = datetime.strptime(f"{date_str} {time_str}", "%Y-%m-%d %H:%M")
    except Exception:
        flash("Bad date/time.", "error")
        return redirect(url_for("interviews"))

    end = start + timedelta(hours=1)
    mode = form.get("Mode","").strip()
    interviewer = form.get("Interviewer Name","").strip()
    
    is_second = (form.get("is_second") == "1")

    if is_second:
        status_value = "Second Interview"
    else:
        status_value = "First Interview"


    iv = _excel_read("Interviews")
    if iv.empty:
        iv = pd.DataFrame(columns=[
            "Candidate ID","Candidate Name","Position",
            "Interview Date","Interview Time",
            "Mode","Location/Link","Meeting Link",
            "Interviewer","ICS Path","Timestamp",
            "Email","Created By","Status"
        ])
    else:
        for col in ["Location/Link","Meeting Link","Status"]:
            if col not in iv.columns:
                iv[col] = ""

    # Onsite behavior
    if mode == "Onsite":
        location_phys = "https://maps.app.goo.gl/WLVM69QkbntDG2Tq7"
        meeting_link = ""
    else:
        location_phys = form.get("Location","").strip()
        meeting_link = form.get("Meeting Link","").strip()

    # ICS + email summary
    if is_second:
        ics_summary = f"Second Interview — {cname} ({position})"
        mail_subject = f"Second Interview Invite – {cname} – {position} – {date_str} {time_str}"
        status_value = "Second Interview"
    else:
        ics_summary = f"Interview — {cname} ({position})"
        mail_subject = f"Interview Invite – {cname} – {position} – {date_str} {time_str}"
        status_value = "First Interview"

    cand_dir = candidate_root(cname, cid)
    ics_path = _make_ics(
        summary=ics_summary,
        description="Interview scheduled via HR System",
        start_dt=start,
        end_dt=end,
        location=location_phys,
        meeting_link=meeting_link,
        attendee_email=(email or None),
        cand_dir=cand_dir,
    )

    new_row = {
        "Candidate ID": cid,
        "Candidate Name": cname,
        "Position": position,
        "Interview Date": date_str,
        "Interview Time": time_str,
        "Mode": mode,
        "Location/Link": location_phys,
        "Meeting Link": meeting_link,
        "Interviewer": interviewer,
        "ICS Path": ics_path,
        "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "Email": email,
        "Created By": current_user().get("username"),
        "Status": status_value,
    }

    # First interview: update existing row if exists, else append
    # Second interview: always append new row
    # detect rows for this candidate
    mask = iv["Candidate ID"].astype(str) == cid

    if not is_second and mask.any():
        # update FIRST interview only
        idx0 = iv.index[mask][0]
        for k, v in new_row.items():
            iv.loc[idx0, k] = v
        msg = "Interview updated."
    else:
        # always create NEW row if this is a second interview
        iv = pd.concat([iv, pd.DataFrame([new_row])], ignore_index=True)
        msg = "Second interview scheduled." if is_second else "Interview saved."


    _excel_write(iv, "Interviews")

    # Optional: keep Candidates sheet status as "Interview"
    cand_df = _excel_read("Candidates")
    if not cand_df.empty and "Candidate ID" in cand_df.columns:
        m = cand_df["Candidate ID"].astype(str) == cid
        if m.any():
            cand_df.loc[m, "Status"] = "Interview"
            cand_df.loc[m, "Last Updated"] = datetime.now()
            _excel_write(cand_df, "Candidates")

    # Email body
    body_lines = [
        f"Dear {cname},",
        "",
        ("You're invited for an interview. Details below:"
         if not is_second else
         "You're invited for a second interview. Details below:"),
        f"- Position: {position}",
        f"- Date: {date_str}",
        f"- Time: {time_str}",
        f"- Location: {location_phys}",
    ]
    if meeting_link:
        body_lines.append(f"- Link: {meeting_link}")
    body_lines += ["", "", "Best regards,"]

    import urllib.parse as _uq
    mailto = ""
    if email:
        body_text = "\r\n".join(body_lines)
        mailto = (
            "mailto:" + _uq.quote(email)
            + "?subject=" + _uq.quote(mail_subject)
            + "&body=" + _uq.quote(body_text)
        )

    short_name = folder_display_name(cname, cid)
    flash(f"{msg} Candidate ID: {cid} (Folder: {short_name})", "success")
    return redirect(url_for("interviews", last_ics=ics_path, mailto=mailto))


@app.get("/interviews/delete/<int:idx>")
def interview_delete(idx: int):
    if not require_role("admin","hr"):
        return redirect(url_for("login"))
    
    iv_df = _excel_read("Interviews")
    if iv_df.empty:
        flash("No interviews to delete.", "error")
        return redirect(url_for("interviews"))

    try:
        if 0 <= idx < len(iv_df):
            cname = iv_df.iloc[idx].get("Candidate Name", "Unknown")
            iv_df = iv_df.drop(iv_df.index[idx])
            _excel_write(iv_df, "Interviews")
            flash(f"Interview for {cname} deleted.", "success")
        else:
            flash("Invalid interview index.", "error")
    except Exception as e:
        flash(f"Error deleting interview: {e}", "error")

    return redirect(url_for("interviews"))

@app.get("/interviews/done_first/<int:idx>")
def interview_done_first(idx):
    iv = _excel_read("Interviews")
    if iv.empty or idx not in iv.index:
        flash("Invalid interview", "error")
        return redirect(url_for("interviews"))

    if iv.loc[idx,"Status"] == "First Interview":
        iv.loc[idx,"Status"] = "First Interview Completed"
        _excel_write(iv,"Interviews")
        flash("First interview completed.","success")

    return redirect(url_for("interviews"))



@app.get("/interviews/undo_first/<int:idx>")
def interview_undo_first(idx):
    iv = _excel_read("Interviews")
    if iv.empty or idx not in iv.index:
        flash("Invalid","error")
        return redirect(url_for("interviews"))

    if iv.loc[idx,"Status"] == "First Interview Completed":
        iv.loc[idx,"Status"] = "First Interview"
        _excel_write(iv,"Interviews")
        flash("Undo successful.","success")

    return redirect(url_for("interviews"))


@app.get("/interviews/request_second/<int:idx>")
def interview_request_second(idx):
    iv = _excel_read("Interviews")
    if iv.empty or idx not in iv.index:
        flash("Invalid","error")
        return redirect(url_for("interviews"))

    cid = str(iv.loc[idx,"Candidate ID"])
    return redirect(url_for("interviews", second_cid=cid))


@app.get("/interviews/done_second/<int:idx>")
def interview_done_second(idx):
    iv = _excel_read("Interviews")
    if iv.empty or idx not in iv.index:
        flash("Invalid","error")
        return redirect(url_for("interviews"))

    if iv.loc[idx,"Status"] == "Second Interview":
        iv.loc[idx,"Status"] = "Second Interview Completed"
        _excel_write(iv,"Interviews")
        flash("Second interview completed.","success")

    return redirect(url_for("interviews"))



@app.get("/interviews/undo_second/<int:idx>")
def interview_undo_second(idx):
    iv = _excel_read("Interviews")
    if iv.empty or idx not in iv.index:
        flash("Invalid","error")
        return redirect(url_for("interviews"))

    if iv.loc[idx,"Status"] == "Second Interview Completed":
        iv.loc[idx,"Status"] = "Second Interview"
        _excel_write(iv,"Interviews")
        flash("Undo successful.","success")

    return redirect(url_for("interviews"))


@app.get("/interviews/done/<int:idx>")
def interview_done(idx: int):
    if not require_role("admin","hr"):
        return redirect(url_for("login"))

    iv_df = _excel_read("Interviews")
    if iv_df.empty or idx < 0 or idx >= len(iv_df):
        flash("Invalid interview.", "error")
        return redirect(url_for("interviews"))

    try:
        # Mark only inside the interview sheet, NOT candidates sheet
        iv_df.loc[idx, "Status"] = "Second Interview"
        _excel_write(iv_df, "Interviews")

        flash("Marked as requiring a second interview.", "success")
    except Exception as e:
        flash(f"Error updating interview status: {e}", "error")

    return redirect(url_for("interviews"))
@app.get("/interviews/undo/<int:idx>")
def interview_undo(idx: int):
    if not require_role("admin","hr"):
        return redirect(url_for("login"))

    iv_df = _excel_read("Interviews")
    if iv_df.empty or idx < 0 or idx >= len(iv_df):
        flash("Invalid interview.", "error")
        return redirect(url_for("interviews"))

    try:
        iv_df.loc[idx, "Status"] = "First Interview"
        _excel_write(iv_df, "Interviews")
        flash("Status reverted.", "success")
    except Exception as e:
        flash(f"Error updating interview status: {e}", "error")

    return redirect(url_for("interviews"))

@app.get("/open-ics")
def open_ics_download():
    if not require_login(): return redirect(url_for("login"))
    path = request.args.get("path","").strip()
    full = os.path.abspath(path)
    if not full.startswith(BASE_DIR) or not os.path.exists(full): return "Not found", 404
    return send_file(full, mimetype="text/calendar", as_attachment=True, download_name=os.path.basename(full))
@app.get("/open-appointment")
def open_appointment():
    path = request.args.get("path")
    if not path or not os.path.exists(path):
        return "Not found", 404

    with open(path, "rb") as f:
        data = f.read()

    return Response(
        data,
        mimetype="text/calendar",
        headers={"Content-Disposition": "inline; filename=invite.ics"}
    )

@app.get("/open-ics-inline")
def open_ics_inline():
    if not require_login(): return redirect(url_for("login"))
    path = request.args.get("path","").strip()
    full = os.path.abspath(path)
    if not full.startswith(BASE_DIR) or not os.path.exists(full): return "Not found", 404
    with open(full, "r", encoding="utf-8") as f: text = f.read()
    return Response(text, mimetype="text/plain")

# -------- Offers (HR/Admin)
@app.get("/offers")
def offers():
    if not require_role("admin","hr"): return redirect(url_for("login"))
    return render_page(OFFERS_HTML, combos=_candidate_combo_list(), selected="", data={}, today=date.today().strftime("%Y-%m-%d"))

@app.post("/offers/load")
def offer_load():
    if not require_role("admin","hr"): return redirect(url_for("login"))
    selected = request.form.get("pick","").strip()
    if not selected or "[" not in selected:
        flash("Select a candidate.", "error"); return redirect(url_for("offers"))
    cid = selected.split("[")[-1].strip("]")
    candidate_data = {}

    df_cand = _excel_read("Candidates")
    df_scr  = _excel_read("Screening_Form")
    c_row = df_cand[df_cand["Candidate ID"].astype(str)==cid]
    if not c_row.empty: candidate_data.update(c_row.iloc[0].to_dict())
    s_row = df_scr[df_scr["Candidate ID"].astype(str)==cid]
    if not s_row.empty: candidate_data.update(s_row.iloc[0].to_dict())

    df_offer = _excel_read("Offer_Details")
    if not df_offer.empty:
        o_row = df_offer[df_offer["Candidate ID"].astype(str)==cid]
        if not o_row.empty:
            for k in ["Basic Salary","Accommodation Allowance","Transportation Allowance", "Monthly Fixed Allowance", "Other Monthly Allowance", "Air Ticket"]:
                if o_row.iloc[0].get(k): candidate_data[k] = str(o_row.iloc[0].get(k))
    
    if not candidate_data.get("Basic Salary") and candidate_data.get("Expected Compensation"):
        candidate_data["Basic Salary"] = str(candidate_data.get("Expected Compensation"))
    
    def _first(keys):
        for k in keys:
            v = candidate_data.get(k)
            if v not in (None, "", "nan", "NaN"): return str(v).strip()
        lower = {str(k).strip().lower(): k for k in candidate_data.keys()}
        for k in keys:
            lk = str(k).strip().lower()
            if lk in lower:
                v = candidate_data.get(lower[lk])
                if v not in (None, "", "nan", "NaN"): return str(v).strip()
        return ""

    if not candidate_data.get("Role Interviewed For"):
        v = _first(["Role Interviewed For","Role","Position","Job Title","Title"])
        if v: candidate_data["Role Interviewed For"] = v
    if not candidate_data.get("Nationality"):
        v = _first(["Nationality","Nationality ","Nationalty","Nationality (as per ID)","Country","الجنسية","Nationality/الجنسية"])
        if v: candidate_data["Nationality"] = v
    if not candidate_data.get("Marital Status"):
        v = _first(["Marital Status","Marital status","Marital","Marital_Status","Married/Single","الحالة الاجتماعية","Status (Marital)"])
        if v: candidate_data["Marital Status"] = v
    if not candidate_data.get("Gov ID / Iqama / Passport #"):
        v = _first(["Gov ID / Iqama / Passport #","Gov ID","National ID","National ID No","ID Number","ID No","Iqama","Iqama No","Iqama Number","Residence Permit","Iqama ID","Passport #","Passport No","Passport Number","Passport"])
        if v: candidate_data["Gov ID / Iqama / Passport #"] = v
    if not candidate_data.get("Candidate Email"):
        v = _first(["Candidate Email","Email","E-mail","Email Address"])
        if v: candidate_data["Candidate Email"] = v
    if not candidate_data.get("Offer Issue Date"):
        from datetime import date as _date
        candidate_data["Offer Issue Date"] = _date.today().strftime("%Y-%m-%d")

    flash("Offer fields loaded.", "success")
    return render_page(OFFERS_HTML, combos=_candidate_combo_list(), selected=selected, data=candidate_data, today=date.today().strftime("%Y-%m-%d"))

@app.post("/offers/generate")
def offer_generate():
    if not require_role("admin","hr"):
        return redirect(url_for("login"))

    selected = request.form.get("selected","").strip()
    if not selected or "[" not in selected:
        flash("Load a candidate first.", "error")
        return redirect(url_for("offers"))

    cid = selected.split("[")[-1].strip("]")
    c_name = selected.split("[")[0].strip()
    payload = {k:(v.strip() if isinstance(v,str) else v) for k,v in request.form.items()}

    # --- LOGIC FOR TEMPLATE SELECTION ---
    nationality = payload.get("Nationality", "").strip().lower()
    location_type = payload.get("Location Type", "Head Office").strip()
    
    # Financial Helpers
    def get_f(key):
        try: return float(payload.get(key, "0") or 0)
        except: return 0.0

    basic = get_f("Basic Salary")
    accom = get_f("Accommodation Allowance")
    trans = get_f("Transportation Allowance")
    fixed = get_f("Monthly Fixed Allowance")
    other = get_f("Other Monthly Allowance")
    air_ticket = payload.get("Air Ticket", "")

    total_pkg = basic + accom + trans + fixed + other

    # Default mappings common to most
    # Common keys: "Position"->B4, "Candidate Name"->B5, "Gov ID"->B6, "Nationality"->B7, "Email"->B8, "Basic"->B9
    base_map = {
        "Role Interviewed For": "B4",
        "Candidate Name": "B5",
        "Gov ID / Iqama / Passport #": "B6",
        "Nationality": "B7",
        "Candidate Email": "B8",
        "Basic Salary": "B9"
    }
    
    # Decision Tree
    template_file = ""
    mapping = base_map.copy()
    
    # 1. SAUDI
    if "saudi" in nationality:
        template_file = TPL_SAUDI
        mapping["Total Package"] = "B12"
        # In Saudi template, we likely just need Basic + Total, maybe Accom/Trans if sheet expects them
        # But prompt says: B9->Basic, B12->Total, B19->Air Ticket.
        payload["Total Package"] = total_pkg
        
    # 2. PHILIPPINES
    elif "filipino" in nationality or "philippines" in nationality:
        if location_type == "Site":
            # Template 4: offer_philippine_site
            template_file = TPL_PHIL_SITE
            mapping["Monthly Fixed Allowance"] = "B10"
            mapping["Other Monthly Allowance"] = "B11"
            mapping["Total Package"] = "B12"
            mapping["Accommodation Allowance"] = "B13"
            payload["Total Package"] = total_pkg
        else:
            # Template 3: offer_philippine_ho
            template_file = TPL_PHIL_HO
            mapping["Total Package"] = "B13"
            payload["Total Package"] = total_pkg
            
    # 3. OTHER FOREIGN
    else:
        if location_type == "Site":
            # Template 2: offer_foreign_site
            template_file = TPL_FOREIGN_SITE
            mapping["Monthly Fixed Allowance"] = "B10"
            mapping["Other Monthly Allowance"] = "B11"
            mapping["Total Package"] = "B12"
            payload["Total Package"] = total_pkg
        else:
            # Template 1: offer_foreign_ho
            template_file = TPL_FOREIGN_HO
            mapping["Total Package"] = "B12"
            payload["Total Package"] = total_pkg

    try:
        result = _generate_offer_doc(cid, c_name, payload, template_file, mapping)
    except Exception as e:
        flash(f"Offer generation failed: {e}", "error")
        return redirect(url_for("offers"))

    if not result:
        flash("Offer generation failed: No output returned.", "error")
        return redirect(url_for("offers"))

    xlsx_path, _ = result

    df = _excel_read("Offer_Details")
    base_cols = [
        "Candidate ID","Candidate Name","Position",
        "Offer Issue Date","Basic Salary","Accommodation Allowance",
        "Transportation Allowance", "Monthly Fixed Allowance", "Other Monthly Allowance", "Air Ticket",
        "Offer Excel Path","Timestamp"
    ]

    if df.empty or "Candidate ID" not in df.columns:
        df = pd.DataFrame(columns=base_cols)

    new_row = {
        "Candidate ID": cid,
        "Candidate Name": c_name,
        "Position": payload.get("Role Interviewed For",""),
        "Offer Issue Date": payload.get("Offer Issue Date", date.today().strftime("%Y-%m-%d")),
        "Basic Salary": payload.get("Basic Salary",""),
        "Accommodation Allowance": payload.get("Accommodation Allowance",""),
        "Transportation Allowance": payload.get("Transportation Allowance",""),
        "Monthly Fixed Allowance": payload.get("Monthly Fixed Allowance",""),
        "Other Monthly Allowance": payload.get("Other Monthly Allowance",""),
        "Air Ticket": payload.get("Air Ticket",""),
        "Offer Excel Path": xlsx_path,
        "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }

    if df.empty or not (df["Candidate ID"].astype(str)==cid).any():
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
    else:
        m = (df["Candidate ID"].astype(str)==cid)
        for k, v in new_row.items():
            if k not in df.columns:
                df[k] = ""
            df.loc[m, k] = v

    _excel_write(df, "Offer_Details")

    cand_df = _excel_read("Candidates")
    if not cand_df.empty and "Candidate ID" in cand_df.columns:
        m = (cand_df["Candidate ID"].astype(str)==cid)
        if m.any():
            cand_df.loc[m, "Status"] = "Offer Issued"
            cand_df.loc[m, "Last Updated"] = datetime.now()
            _excel_write(cand_df, "Candidates")

    flash(f"Offer created ({template_file}) for {c_name}", "success")
    return redirect(url_for("offers"))


# -------- Candidates (HR/Admin/Requestor)
@app.get("/candidates")
def candidates():
    if not require_role("admin","hr","requestor"): 
        return redirect(url_for("login"))

    q = {k: v for k, v in request.args.items()}

    u = current_user()
    req_user = u["username"] if u["role"] == "requestor" else None

    rows, _ = _candidate_rows(q, requestor_user=req_user)

    return render_page(
        CANDIDATES_HTML,
        rows=rows,
        q=q,
        statuses=CAND_STATUS,
        cand=None,
        cand_id=None,
        cand_name=None,
        checklist=[],
        YESNO=YESNO,
        meta={},
        req_actions=REQ_ACTIONS
    )


@app.get("/candidates/<cand_id>")
def candidate_detail(cand_id: str):
    if not require_role("admin","hr","requestor"): return redirect(url_for("login"))
    q = {k:v for k,v in request.args.items()}
    
    cand_df = _excel_read("Candidates")
    row = cand_df[cand_df["Candidate ID"].astype(str)==cand_id]
    if row.empty:
        flash("Candidate not found.", "error"); return redirect(url_for("candidates"))
    
    # Strict Requestor Check
    u = current_user()
    if u and u.get("role")=="requestor":
        if str(row.iloc[0].get("Requestor Username","")) != u["username"]:
            flash("Access denied.", "error"); return redirect(url_for("candidates"))
            
    cand_name = row.iloc[0].get("Candidate Name","")
    
    # ---- FETCH CV FROM SCREENING FORM (Source of Truth) ----
    # This fixes the bug where imported candidates had CV in Screening but not in Candidates sheet
        # ---- FETCH CV FROM SCREENING FORM (Source of Truth) ----
    # This fixes the bug where imported candidates had CV in Screening but not in Candidates sheet
    sf_df = _excel_read("Screening_Form")
    cv_path = ""
    if not sf_df.empty:
        sf_row = sf_df[sf_df["Candidate ID"].astype(str)==cand_id]
        if not sf_row.empty:
            cv_path = str(sf_row.iloc[0].get("CV File Path","") or "").strip()
            if cv_path.lower() == "nan": cv_path = ""

    # Fallback to Candidates sheet if Screening is empty but Candidates has it
    if not cv_path:
        cv_path = str(row.iloc[0].get("CV File Path","") or "").strip()
        if cv_path.lower() == "nan": cv_path = ""

    # Final CV path sanity check – hide if file is not actually present/allowed
    if cv_path:
        _cv_full = os.path.abspath(cv_path)
        if not (_cv_full.startswith(BASE_DIR) and os.path.exists(_cv_full)):
            cv_path = ""

    # Load metadata
    meta = {
        "Status": row.iloc[0].get("Status","Screening"),
        "Next Action": row.iloc[0].get("Next Action",""),
        "Notes": row.iloc[0].get("Notes",""),
        "Requestor Comments": row.iloc[0].get("Requestor Comments",""),
        "Requestor Action": row.iloc[0].get("Requestor Action",""),
        "Suggested Interview Date": row.iloc[0].get("Suggested Interview Date",""),
        "Suggested Interview Time": row.iloc[0].get("Suggested Interview Time",""),
        # Use the resolved CV path
        "CV File Path": cv_path,
    }


    # checklist
    sl_df = _excel_read("Shortlist_Request")
    checklist=[]
    if sl_df.empty or sl_df[sl_df["Candidate ID"].astype(str)==cand_id].empty:
        defaults = ["CV/Resume","Passport/Iqama Copy","Education Certificate","Experience Letters","Requestor Assessment (Internal)"]
        for it in defaults:
            checklist.append({"Item":it,"Received (Yes/No)":"No","Notes":"","Mapped File Path":""})
    else:
        rows_sl = sl_df[sl_df["Candidate ID"].astype(str)==cand_id]
        for _, r in rows_sl.iterrows():
            mapped = str(r.get("Mapped File Path","") or "").strip()
            if mapped.lower() == "nan": mapped = ""
            if mapped:
                _m_full = os.path.abspath(mapped)
                if not (_m_full.startswith(BASE_DIR) and os.path.exists(_m_full)):
                    mapped = ""
            checklist.append({
                "Item": r["Item"],
                "Received (Yes/No)": normalize_choice(r.get("Received (Yes/No)",""), YESNO),
                "Notes": r.get("Notes",""),
                "Mapped File Path": mapped
            })

    # Re-fetch list for sidebar (filtered)
    req_user = u["username"] if u["role"] == "requestor" else None
    rows, _ = _candidate_rows(q, requestor_user=req_user)

    return render_page(CANDIDATES_HTML,
        rows=rows, q=q, statuses=CAND_STATUS,
        cand=True, cand_id=cand_id, cand_name=cand_name,
        checklist=checklist, YESNO=YESNO, meta=meta, req_actions=REQ_ACTIONS
    )

@app.post("/candidates/meta/save/<cand_id>")
def candidate_meta_save(cand_id: str):
    if not require_role("admin","hr"): return redirect(url_for("login"))
    cand_df = _excel_read("Candidates")
    if cand_df.empty:
        flash("Candidates sheet empty.", "error"); return redirect(url_for("candidates"))
    m = (cand_df["Candidate ID"].astype(str)==cand_id)
    if not m.any():
        flash("Candidate not found.", "error"); return redirect(url_for("candidates"))
    cand_df.loc[m, "Status"] = request.form.get("Status","Screening")
    cand_df.loc[m, "Next Action"] = request.form.get("Next Action","")
    cand_df.loc[m, "Notes"] = request.form.get("Notes","")
    cand_df.loc[m, "Last Updated"] = datetime.now()
    _excel_write(cand_df, "Candidates")
    flash("Candidate updated.", "success")
    return redirect(url_for("candidate_detail", cand_id=cand_id))

@app.post("/candidates/comment/<cand_id>")
def candidate_comment(cand_id: str):
    if not require_role("requestor","admin","hr"): return redirect(url_for("login"))
    u = current_user()
    cand_df = _excel_read("Candidates")
    if cand_df.empty:
        flash("Candidates sheet empty.", "error"); return redirect(url_for("candidates"))
    m = (cand_df["Candidate ID"].astype(str)==cand_id)
    if not m.any():
        flash("Candidate not found.", "error"); return redirect(url_for("candidates"))
    
    if u and u.get("role")=="requestor":
        if str(cand_df.loc[m].iloc[0].get("Requestor Username","")) != u["username"]:
            flash("Access denied.", "error"); return redirect(url_for("candidates"))

    cand_df.loc[m, "Requestor Comments"] = request.form.get("Requestor Comments","")
    cand_df.loc[m, "Requestor Action"] = request.form.get("Requestor Action","")
    cand_df.loc[m, "Suggested Interview Date"] = request.form.get("Suggested Interview Date","")
    cand_df.loc[m, "Suggested Interview Time"] = request.form.get("Suggested Interview Time","")
    
    cand_df.loc[m, "Last Updated"] = datetime.now()
    _excel_write(cand_df, "Candidates")
    flash("Request/Comment sent to HR.", "success")
    return redirect(url_for("candidate_detail", cand_id=cand_id))

@app.get("/candidates/delete/<cand_id>")
def candidate_delete_all(cand_id: str):
    if not require_role("admin","hr"): return redirect(url_for("login"))
    sheets = ["Candidates","Screening_Form","Shortlist_Request","Interviews","Offer_Details"]
    for s in sheets:
        df = _excel_read(s)
        if not df.empty:
            df = df[df["Candidate ID"].astype(str)!=str(cand_id)]
            _excel_write(df, s)
    for entry in os.listdir(ATTACH_DIR):
        if entry.endswith("_"+cand_id) and os.path.isdir(os.path.join(ATTACH_DIR, entry)):
            try: shutil.rmtree(os.path.join(ATTACH_DIR, entry))
            except Exception: pass
    flash(f"Candidate deleted: {cand_id}", "success")
    return redirect(url_for("candidates"))

@app.post("/shortlist/save/<cand_id>")
def shortlist_save(cand_id: str):
    if not require_role("admin","hr"): return redirect(url_for("login"))
    rows = int(request.form.get("rows","0") or 0)
    new_item = (request.form.get("new_item") or "").strip()

    cdf = _excel_read("Candidates")
    r = cdf[cdf["Candidate ID"].astype(str)==cand_id]
    if r.empty:
        flash("Candidate not found.", "error"); return redirect(url_for("candidates"))
    cand_name = r.iloc[0].get("Candidate Name","")
    dst_dir = candidate_attach_dir(cand_name, cand_id)

    prev_df = _excel_read("Shortlist_Request")
    prev_for_cand = None
    if not prev_df.empty:
        prev_for_cand = prev_df[prev_df["Candidate ID"].astype(str)==cand_id]

    all_rows=[]
    for i in range(1, rows+1):
        item = request.form.get(f"item_{i}","").strip()
        recv_req = request.form.get(f"recv_{i}","").strip()
        note = request.form.get(f"note_{i}","").strip()
        mapped = ""
        uploaded = False

        f = request.files.get(f"map_{i}")
        if f and f.filename:
            fn = secure_filename(f.filename)
            target = os.path.join(dst_dir, fn)
            f.save(target)
            mapped = target
            uploaded = True
        else:
            if prev_for_cand is not None and not prev_for_cand.empty and item:
                prev = prev_for_cand[prev_for_cand["Item"].astype(str)==item]
                if not prev.empty:
                    mapped = str(prev.iloc[0].get("Mapped File Path","") or "").strip()
                    if mapped.lower() == "nan": mapped = ""

        recv = "Yes" if uploaded else normalize_choice(recv_req, YESNO)

        if item:
            all_rows.append({
                "Candidate ID": cand_id, "Candidate Name": cand_name,
                "Item": item, "Received (Yes/No)": recv, "Notes": note,
                "Mapped File Path": mapped, "Timestamp": datetime.now()
            })

    if new_item:
        all_rows.append({
            "Candidate ID": cand_id, "Candidate Name": cand_name,
            "Item": new_item, "Received (Yes/No)": "No", "Notes": "",
            "Mapped File Path": "", "Timestamp": datetime.now()
        })

    if not all_rows:
        flash("Checklist empty.", "error"); return redirect(url_for("candidate_detail", cand_id=cand_id))

    sl_df = _excel_read("Shortlist_Request")
    if not sl_df.empty:
        sl_df = sl_df[sl_df["Candidate ID"].astype(str)!=cand_id]
    new_df = pd.DataFrame(all_rows)
    sl_df = pd.concat([sl_df, new_df], ignore_index=True)
    _excel_write(sl_df, "Shortlist_Request")
    flash("Shortlist saved.", "success")
    return redirect(url_for("candidate_detail", cand_id=cand_id))

@app.get("/shortlist/delete/<cand_id>/<int:idx>")
def shortlist_delete(cand_id: str, idx: int):
    if not require_role("admin","hr"): return redirect(url_for("login"))
    sl_df = _excel_read("Shortlist_Request")
    if sl_df.empty:
        flash("Nothing to delete.", "error"); return redirect(url_for("candidate_detail", cand_id=cand_id))
    rows = sl_df[sl_df["Candidate ID"].astype(str)==cand_id]
    if rows.empty or idx<0 or idx>=len(rows):
        flash("Nothing to delete.", "error"); return redirect(url_for("candidate_detail", cand_id=cand_id))
    target_item = rows.iloc[idx]["Item"]
    sl_df = sl_df[~((sl_df["Candidate ID"].astype(str)==cand_id) & (sl_df["Item"].astype(str)==target_item))]
    _excel_write(sl_df, "Shortlist_Request")
    flash("Item deleted.", "success")
    return redirect(url_for("candidate_detail", cand_id=cand_id))
@app.get("/shortlist/remove-file/<cand_id>/<int:idx>")
def shortlist_remove_file(cand_id: str, idx: int):
    if not require_role("admin","hr"):
        return redirect(url_for("login"))

    sl_df = _excel_read("Shortlist_Request")
    if sl_df.empty:
        flash("Nothing to update.", "error")
        return redirect(url_for("candidate_detail", cand_id=cand_id))

    rows = sl_df[sl_df["Candidate ID"].astype(str) == cand_id]
    if rows.empty or idx < 0 or idx >= len(rows):
        flash("Invalid index.", "error")
        return redirect(url_for("candidate_detail", cand_id=cand_id))

    item_name = rows.iloc[idx]["Item"]

    # Remove ONLY the file path (not the row)
    sl_df.loc[
        (sl_df["Candidate ID"].astype(str) == cand_id) &
        (sl_df["Item"] == item_name),
        "Mapped File Path"
    ] = ""

    _excel_write(sl_df, "Shortlist_Request")

    flash("Attachment removed successfully.", "success")
    return redirect(url_for("candidate_detail", cand_id=cand_id))

# -------- File openers
@app.get("/open")
def open_inline():
    if not require_login():
        return redirect(url_for("login"))

    path = request.args.get("path", "").strip()

    # prevent empty / invalid / non-existing paths
    if not path or path.lower() in ["", "none", "nan"]:
        return "File not available", 404

    full = os.path.abspath(path)

    if not full.startswith(BASE_DIR):
        return "Invalid file path", 404

    if not os.path.exists(full):
        return "File not found", 404

    mime, _ = mimetypes.guess_type(full)
    return send_file(full, mimetype=mime or "application/octet-stream",
                     as_attachment=False,
                     download_name=os.path.basename(full))

# ---------- UTIL HELPERS (data)
def _list_requestors():
    users = _load_users()
    out=[]
    for u in users.values():
        if u.get("role")=="requestor":
            out.append(type("U",(object,),u))
    return sorted(out, key=lambda x: (x.name or x.username).lower())

def _screening_picker(filter_user: str = None) -> List[str]:
    df = _excel_read("Screening_Form")
    items = []
    if not df.empty:
        for _, r in df.iterrows():
            # If filter_user is set (Requestor mode), only show their candidates
            if filter_user and str(r.get("Requestor Username","")) != filter_user:
                continue
            items.append(f"{r.get('Candidate Name','')} [{r.get('Candidate ID','')}]")
    return sorted(items)

def _candidate_combo_list() -> List[str]:
    items=[]
    df = _excel_read("Candidates")
    for _, r in df.iterrows():
        nm = str(r.get('Candidate Name','')).strip()
        cid = str(r.get('Candidate ID','')).strip()
        if nm and cid:
            items.append(f"{nm} [{cid}]")
    return sorted(items)

def _load_screening_row(cand_id: str) -> Dict[str, Any]:
    sf = _excel_read("Screening_Form")
    if sf.empty: return {}
    row = sf[sf["Candidate ID"].astype(str)==str(cand_id)]
    return {} if row.empty else row.iloc[0].to_dict()

def _interview_cand_map() -> Dict[str, Dict[str,str]]:
    cand = _excel_read("Candidates")
    scr  = _excel_read("Screening_Form")
    out: Dict[str, Dict[str,str]] = {}
    scr_idx = {}
    for _, r in scr.iterrows():
        scr_idx[str(r.get("Candidate ID",""))] = r
    for _, r in cand.iterrows():
        cid = str(r.get("Candidate ID","")).strip()
        nm  = str(r.get("Candidate Name","")).strip()
        if not cid or not nm: continue
        key = f"{nm} [{cid}]"
        email = ""
        role  = str(r.get("Role","")).strip()
        if cid in scr_idx:
            srow = scr_idx[cid]
            email = str(srow.get("Candidate Email","") or "").strip()
            if not role:
                role = str(srow.get("Role Interviewed For","") or "").strip()
        out[key] = {"email": email, "role": role, "cid": cid}
    return out

def _candidate_rows(filters: Dict[str,str], requestor_user: str = None) -> Tuple[List[Dict[str,Any]], List[Dict[str,Any]]]:
    df = _excel_read("Candidates")
    view = df.copy()

    # Requestor filter
    if requestor_user:
        view = view[view["Requestor Username"] == requestor_user]

    # Existing filters
    nm = (filters.get("name","") or "").strip().lower()
    rl = (filters.get("role","") or "").strip().lower()
    st = filters.get("status","All")

    # NEW: Filter by Candidate ID / Gov ID / Iqama
    sid = (filters.get("search_id","") or "").strip().lower()

    if nm:
        view = view[view["Candidate Name"].astype(str).str.lower().str.contains(nm, na=False)]

    if rl:
        view = view[view["Role"].astype(str).str.lower().str.contains(rl, na=False)]

    if st and st != "All":
        view = view[view["Status"].astype(str) == st]

    # Apply ID/Iqama filter on BOTH sheets (Candidates + Screening_Form)
    if sid:
        sf = _excel_read("Screening_Form")
        if not sf.empty:
            sf_match = sf[
                sf["Candidate ID"].astype(str).str.lower().str.contains(sid, na=False)
                | sf["Gov ID / Iqama / Passport #"].astype(str).str.lower().str.contains(sid, na=False)
            ]
            valid_ids = set(sf_match["Candidate ID"].astype(str))
            view = view[view["Candidate ID"].astype(str).isin(valid_ids)]

    rows = view.to_dict("records")
    return rows, df.to_dict("records")

def _dashboard_data():
    df = _excel_read("Candidates")
    totals = {"total":0,"new_week":0,"with_cv":0,"shortlisted":0,"interview":0,"offer_issued":0,"offer_accepted":0,"on_hold":0,"rejected":0}
    status_counts = {s:0 for s in CAND_STATUS}
    recent = []
    if not df.empty:
        now = datetime.now()
        if "Last Updated" in df.columns:
            df["Last Updated"] = pd.to_datetime(df["Last Updated"], errors="coerce")
        else:
            df["Last Updated"] = pd.NaT
        for _, r in df.iterrows():
            st = str(r.get("Status","Other") or "Other")
            if st not in status_counts: status_counts[st] = 0
            status_counts[st] += 1
            cvp = str(r.get("CV File Path","") or "")
            if cvp.strip(): totals["with_cv"] += 1
            lu = r.get("Last Updated")
            if pd.notna(lu) and (now - lu).days <= 7:
                totals["new_week"] += 1
        totals["total"] = int(len(df.index))
        totals["shortlisted"]    = status_counts.get("Shortlist",0)
        totals["interview"]      = status_counts.get("Interview",0) + status_counts.get("Second Interview",0)
        totals["offer_issued"]   = status_counts.get("Offer Issued",0)
        totals["offer_accepted"] = status_counts.get("Offer Accepted",0)
        totals["on_hold"]        = status_counts.get("On Hold",0)
        totals["rejected"]       = status_counts.get("Rejected",0)
        recent_df = df.sort_values(by="Last Updated", ascending=False, na_position="last").head(10)
        out=[]
        for _, r in recent_df.iterrows():
            rr = r.to_dict()
            lu = rr.get("Last Updated")
            rr["Last Updated"] = lu.strftime("%Y-%m-%d %H:%M") if pd.notna(lu) else ""
            out.append(rr)
        recent = out
    status_rows = [{"label": s, "count": status_counts.get(s,0)} for s in CAND_STATUS]
    return totals, status_rows, recent

# ---------- INLINE BASE LOADER
app.jinja_env.globals["base"] = BASE_HTML
app.jinja_loader = type("InlineLoader",(object,),{
    "get_source": lambda self, env, template: (
        BASE_HTML if template=="base.html" else
        LOGIN_HTML if template=="login.html" else
        HOME_HTML if template=="home.html" else
        SCREENING_HTML if template=="screening.html" else
        INTERVIEWS_HTML if template=="interviews.html" else
        OFFERS_HTML if template=="offers.html" else
        CANDIDATES_HTML if template=="candidates.html" else
        USERS_HTML if template=="users.html" else
        "", template, lambda: True
    )
})()

# ---------- ICS BUILDER
def _make_ics(summary: str, description: str, start_dt: datetime, end_dt: datetime, location: str, meeting_link: str, attendee_email: Optional[str], cand_dir: str) -> str:
    def fmt(dt): return dt.astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    uid = uuid.uuid4().hex
    
    # Combine location and link for visibility
    full_location = location
    if meeting_link:
        full_location = f"{location} (Link in desc)" if location else "Online"
        description = f"Meeting Link: {meeting_link}\n\n{description}"
    
    lines = [
        "BEGIN:VCALENDAR","VERSION:2.0","PRODID:-//QNCS HR System//Interview Invite//EN",
        "BEGIN:VEVENT",
        f"UID:{uid}", f"DTSTAMP:{fmt(datetime.now(timezone.utc))}",
        f"DTSTART:{fmt(start_dt)}", f"DTEND:{fmt(end_dt)}",
        f"SUMMARY:{summary}", f"DESCRIPTION:{description}", f"LOCATION:{full_location}",
    ]
    if attendee_email:
        lines.append(f"ATTENDEE;CN=Candidate;ROLE=REQ-PARTICIPANT:MAILTO:{attendee_email}")
    lines += ["END:VEVENT","END:VCALENDAR"]
    os.makedirs(cand_dir, exist_ok=True)
    path = os.path.join(cand_dir, f"{uid}.ics")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\r\n".join(lines))
    return path

if __name__ == "__main__":
    app.run(debug=True)